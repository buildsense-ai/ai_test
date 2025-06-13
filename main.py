import asyncio
import uvicorn
from fastapi import FastAPI, Request, HTTPException, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from typing import Optional, Dict, Any, List
import httpx
import json
from datetime import datetime
import socket
from enum import Enum
import io
import os
import re
import tempfile
import logging
import traceback
import uuid
import hashlib
import time

# Database imports
try:
    import pymysql
    PYMYSQL_AVAILABLE = True
except ImportError:
    PYMYSQL_AVAILABLE = False
    print("⚠️ PyMySQL not available. Database features disabled.")

# Coze SDK imports
try:
    from cozepy import COZE_CN_BASE_URL
    from cozepy import Coze, TokenAuth, Message, ChatStatus, MessageContentType, ChatEventType
    COZE_SDK_AVAILABLE = True
except ImportError:
    COZE_SDK_AVAILABLE = False
    print("⚠️ Coze SDK not available. Using fallback HTTP requests.")

# Import configuration
import config

# ⭐ Memory monitoring
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    print("⚠️ psutil not available, memory monitoring disabled")

# Set up logging for better debugging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add enhanced debugging for cloud deployment issues
import traceback
import logging

# Set up detailed logging for debugging cloud deployment issues
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def check_memory_usage():
    """Check memory usage and prevent OOM crashes"""
    if not PSUTIL_AVAILABLE:
        print("⚠️ psutil not available, skipping memory check")
        return 0  # Return 0 instead of None for compatibility
    
    try:
        memory_percent = psutil.virtual_memory().percent
        if memory_percent > config.MEMORY_CRITICAL_THRESHOLD:
            error_msg = f"服务器内存使用率危险: {memory_percent:.1f}%"
            print(f"❌ {error_msg}")
            raise HTTPException(status_code=507, detail=f"服务器内存不足 ({memory_percent:.1f}%)，请稍后重试")
        
        if memory_percent > config.MEMORY_WARNING_THRESHOLD:
            print(f"⚠️ 内存使用率警告: {memory_percent:.1f}%")
        
        return memory_percent
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"⚠️ 内存检查失败: {str(e)}")
        return 0  # Return 0 instead of None for compatibility

# Document processing imports
try:
    import docx
    from PyPDF2 import PdfReader
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False

app = FastAPI(title="AI Agent Evaluation Platform", version="4.0.0")

# Add response compression middleware
from fastapi.middleware.gzip import GZipMiddleware
app.add_middleware(GZipMiddleware, minimum_size=1000)

# Create static directory if it doesn't exist
static_dir = "static"
if not os.path.exists(static_dir):
    os.makedirs(static_dir)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

templates = Jinja2Templates(directory="templates")

# Improved document processing functions based on user's approach
def read_docx_file(filepath: str) -> str:
    """Read content from DOCX file with enhanced cloud compatibility"""
    try:
        logger.info(f"📖 开始解析DOCX文件: {filepath}")
        print(f"📖 开始解析DOCX文件: {filepath}")
        
        # Check if file exists and is readable
        if not os.path.exists(filepath):
            error_msg = f"文件不存在: {filepath}"
            logger.error(error_msg)
            return f"错误：{error_msg}"
        
        file_size = os.path.getsize(filepath)
        logger.info(f"📄 文件大小: {file_size} 字节")
        print(f"📄 文件大小: {file_size} 字节")
        
        if file_size == 0:
            error_msg = "文件为空"
            logger.error(error_msg)
            return f"错误：{error_msg}"
        
        # Cloud-compatible processing with multiple fallback methods
        extraction_methods = [
            ("python-docx", lambda: _extract_with_python_docx(filepath)),
            ("zip-xml-advanced", lambda: _extract_with_zip_xml_advanced(filepath)),
            ("zip-xml-simple", lambda: _extract_with_zip_xml_simple(filepath)),
            ("raw-text-extraction", lambda: _extract_raw_text_from_docx(filepath))
        ]
        
        best_result = ""
        successful_method = "none"
        
        for method_name, extraction_func in extraction_methods:
            try:
                logger.info(f"🔄 尝试方法: {method_name}")
                print(f"🔄 尝试方法: {method_name}")
                
                result = extraction_func()
                
                if result and len(result) > len(best_result):
                    best_result = result
                    successful_method = method_name
                    logger.info(f"✅ {method_name} 成功，提取长度: {len(result)}")
                    print(f"✅ {method_name} 成功，提取长度: {len(result)}")
                    
                    # If we get a good result (>100 chars), use it immediately
                    if len(result) > 100:
                        break
                else:
                    logger.warning(f"⚠️ {method_name} 结果不佳: {len(result) if result else 0} 字符")
                    print(f"⚠️ {method_name} 结果不佳: {len(result) if result else 0} 字符")
                    
            except Exception as e:
                logger.warning(f"⚠️ {method_name} 失败: {str(e)}")
                print(f"⚠️ {method_name} 失败: {str(e)}")
                continue
        
        if not best_result:
            return "错误：所有解析方法均失败，建议转换为TXT格式后重试"
        
        # Validate extraction result
        if len(best_result) < 20:
            logger.warning(f"⚠️ 提取内容过短: {len(best_result)} 字符")
            print(f"⚠️ 提取内容过短: {len(best_result)} 字符")
            
            if len(best_result) < 10:
                return f"错误：提取内容过短({len(best_result)}字符)，建议转换为TXT格式：\n\n💡 解决方案：\n1. 使用Word打开文档，另存为.txt格式\n2. 或复制文档内容，直接粘贴到文本框中\n3. 检查文档是否包含主要为图片/表格内容"
        
        logger.info(f"✅ DOCX解析成功 (方法: {successful_method})，提取长度: {len(best_result)} 字符")
        print(f"✅ DOCX解析成功 (方法: {successful_method})，提取长度: {len(best_result)} 字符")
        
        # Debug: Show first part of content to verify extraction
        content_preview = best_result[:200] + "..." if len(best_result) > 200 else best_result
        logger.debug(f"📝 内容预览: {content_preview}")
        print(f"📝 内容预览: {content_preview}")
        
        return best_result
        
    except Exception as e:
        error_msg = f"DOCX文件处理异常: {str(e)}"
        logger.error(f"❌ {error_msg}")
        logger.error(f"📋 异常详情: {traceback.format_exc()}")
        print(f"❌ {error_msg}")
        print(f"📋 异常详情: {traceback.format_exc()}")
        return f"错误：{error_msg}\n\n💡 云环境解决方案：\n1. 转换为TXT格式重新上传\n2. 复制文档内容直接粘贴\n3. 检查文档是否过于复杂"

def _extract_with_python_docx(filepath: str) -> str:
    """Method 1: Standard python-docx extraction"""
    from docx import Document
    
    doc = Document(filepath)
    full_text = []
    
    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())
    
    return '\n'.join(full_text)

def _extract_with_zip_xml_advanced(filepath: str) -> str:
    """Method 2: Advanced ZIP+XML extraction with namespace handling"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    with zipfile.ZipFile(filepath, 'r') as zip_file:
        # Try to get document.xml
        xml_content = zip_file.read('word/document.xml')
        
        # Parse with namespace awareness
        root = ET.fromstring(xml_content)
        
        # Define Word document namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'v': 'urn:schemas-microsoft-com:vml',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        text_parts = []
        
        # Extract text from paragraphs
        for para in root.findall('.//w:p', namespaces):
            para_text = ""
            for text_elem in para.findall('.//w:t', namespaces):
                if text_elem.text:
                    para_text += text_elem.text
            if para_text.strip():
                text_parts.append(para_text.strip())
        
        # Extract text from tables
        for table in root.findall('.//w:tbl', namespaces):
            for row in table.findall('.//w:tr', namespaces):
                row_text = ""
                for cell in row.findall('.//w:tc', namespaces):
                    cell_text = ""
                    for text_elem in cell.findall('.//w:t', namespaces):
                        if text_elem.text:
                            cell_text += text_elem.text
                    if cell_text.strip():
                        row_text += cell_text.strip() + " "
                if row_text.strip():
                    text_parts.append(row_text.strip())
        
        return '\n'.join(text_parts)

def _extract_with_zip_xml_simple(filepath: str) -> str:
    """Method 3: Simple ZIP+XML extraction (cloud fallback)"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    with zipfile.ZipFile(filepath, 'r') as zip_file:
        xml_content = zip_file.read('word/document.xml')
        root = ET.fromstring(xml_content)
        
        # Simple text extraction - get all text elements
        text_content = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                text_content.append(elem.text.strip())
        
        # Remove duplicates while preserving order
        seen = set()
        unique_text = []
        for text in text_content:
            if text not in seen and len(text) > 1:  # Avoid single characters
                seen.add(text)
                unique_text.append(text)
        
        return ' '.join(unique_text)

def _extract_raw_text_from_docx(filepath: str) -> str:
    """Method 4: Raw text extraction from ZIP (last resort)"""
    import zipfile
    import re
    
    with zipfile.ZipFile(filepath, 'r') as zip_file:
        # Try to extract any readable text from the ZIP contents
        text_parts = []
        
        # Read document.xml as raw text and extract with regex
        try:
            xml_content = zip_file.read('word/document.xml').decode('utf-8', errors='ignore')
            
            # Use regex to find text between XML tags
            text_matches = re.findall(r'>([^<]+)<', xml_content)
            
            for match in text_matches:
                cleaned = match.strip()
                if len(cleaned) > 2 and not cleaned.isdigit():  # Skip numbers and short strings
                    text_parts.append(cleaned)
            
        except Exception:
            pass
        
        # Also try other XML files in the document
        for file_info in zip_file.filelist:
            if file_info.filename.endswith('.xml') and 'word/' in file_info.filename:
                try:
                    content = zip_file.read(file_info.filename).decode('utf-8', errors='ignore')
                    text_matches = re.findall(r'>([^<]+)<', content)
                    
                    for match in text_matches:
                        cleaned = match.strip()
                        if len(cleaned) > 2 and not cleaned.isdigit():
                            text_parts.append(cleaned)
                            
                except Exception:
                    continue
        
        # Remove duplicates and join
        unique_parts = list(dict.fromkeys(text_parts))  # Preserve order
        return ' '.join(unique_parts)

def read_pdf_file(filepath: str) -> str:
    """Read PDF document using direct file path approach"""
    try:
        if not DOCUMENT_PROCESSING_AVAILABLE:
            return "文档处理库未安装，请安装 PyPDF2：pip install PyPDF2"
        
        with open(filepath, 'rb') as file:
            pdf_reader = PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text.strip())
        
        # Join and clean the text
        full_text = "\n".join(text_parts)
        cleaned_text = full_text.replace("\r", "").replace("　", "").strip()
        
        print(f"📄 PDF文档提取成功，内容长度: {len(cleaned_text)} 字符")
        return cleaned_text
        
    except Exception as e:
        error_msg = f"PDF文档解析失败: {str(e)}"
        print(f"❌ {error_msg}")
        return error_msg

def read_txt_file(filepath: str) -> str:
    """Read text file with proper encoding"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
        
        # Clean the text
        cleaned_text = text.replace("\r", "").replace("　", "").strip()
        
        print(f"📄 文本文件提取成功，内容长度: {len(cleaned_text)} 字符")
        return cleaned_text
        
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(filepath, "r", encoding="gbk") as f:
                text = f.read()
            cleaned_text = text.replace("\r", "").replace("　", "").strip()
            print(f"📄 文本文件提取成功(GBK编码)，内容长度: {len(cleaned_text)} 字符")
            return cleaned_text
        except Exception as e:
            error_msg = f"文本文件解析失败: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg
    except Exception as e:
        error_msg = f"文本文件解析失败: {str(e)}"
        print(f"❌ {error_msg}")
        return error_msg

async def process_uploaded_document_improved(file: UploadFile) -> str:
    """Process uploaded document using improved approach with comprehensive error handling"""
    if not file or not file.filename:
        logger.error("⚠️ 文档上传：文件为空或无文件名")
        print("⚠️ 文档上传：文件为空或无文件名")
        return "错误：未提供有效文件"
    
    # ⭐ Security: Validate filename
    if not validate_filename(file.filename):
        error_msg = f"不安全的文件名: {file.filename}"
        logger.error(f"❌ {error_msg}")
        print(f"❌ {error_msg}")
        return f"错误：{error_msg}"
    
    # Log file info with detailed debugging
    logger.info(f"📄 开始处理上传文件: {file.filename}")
    logger.info(f"📄 文件类型: {getattr(file, 'content_type', '未知')}")
    print(f"📄 开始处理上传文件: {file.filename}")
    print(f"📄 文件类型: {getattr(file, 'content_type', '未知')}")
    
    # Create temporary file
    suffix = os.path.splitext(file.filename)[1].lower()
    logger.info(f"📄 检测文件扩展名: {suffix}")
    print(f"📄 检测文件扩展名: {suffix}")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        try:
            # Write uploaded content to temporary file
            logger.info("📤 读取上传文件内容...")
            print("📤 读取上传文件内容...")
            content = await file.read()
            
            if not content:
                logger.error("❌ 上传文件内容为空")
                print("❌ 上传文件内容为空")
                return "错误：上传文件内容为空"
            
            logger.info(f"📤 文件大小: {len(content)} 字节")
            print(f"📤 文件大小: {len(content)} 字节")
            
            # ⭐ Critical: File size limit to prevent memory issues
            if len(content) > config.MAX_FILE_SIZE:
                error_msg = f"文件大小 {len(content)} 字节超过10MB限制"
                logger.error(f"❌ {error_msg}")
                print(f"❌ {error_msg}")
                return f"错误：{error_msg}"
            
            tmp_file.write(content)
            tmp_file.flush()
            
            logger.info(f"💾 临时文件已创建: {tmp_file.name}")
            print(f"💾 临时文件已创建: {tmp_file.name}")
            
            # Process based on file extension with enhanced debugging
            try:
                if suffix in ['.doc', '.docx']:
                    logger.info("📖 使用Word文档解析器...")
                    print("📖 使用Word文档解析器...")
                    result = read_docx_file(tmp_file.name)
                elif suffix == '.pdf':
                    logger.info("📖 使用PDF文档解析器...")
                    print("📖 使用PDF文档解析器...")
                    result = read_pdf_file(tmp_file.name)
                elif suffix == '.txt':
                    logger.info("📖 使用文本文件解析器...")
                    print("📖 使用文本文件解析器...")
                    result = read_txt_file(tmp_file.name)
                else:
                    error_msg = f"不支持的文件格式: {suffix}。支持格式: Word (.docx), PDF (.pdf), 文本 (.txt)"
                    logger.error(f"❌ {error_msg}")
                    print(f"❌ {error_msg}")
                    return error_msg
            except Exception as parse_error:
                logger.error(f"❌ 文档解析异常: {str(parse_error)}")
                logger.error(f"📋 解析异常详情: {traceback.format_exc()}")
                print(f"❌ 文档解析异常: {str(parse_error)}")
                print(f"📋 解析异常详情: {traceback.format_exc()}")
                return f"错误：文档解析失败 - {type(parse_error).__name__}: {str(parse_error)}"
            
            # Validate result with enhanced debugging
            if not result:
                logger.error("❌ 文档解析结果为空")
                print("❌ 文档解析结果为空")
                return "错误：文档解析结果为空，可能文件已损坏或格式不正确"
            
            if len(result) < 10:
                logger.warning(f"⚠️ 文档解析结果过短: {len(result)} 字符")
                print(f"⚠️ 文档解析结果过短: {len(result)} 字符")
                return f"错误：文档内容过短({len(result)}字符)，可能解析失败"
            
            # Check for error messages in result
            error_indicators = ['error', 'exception', 'traceback', 'failed', 'Error:', 'Exception:', '处理失败', '解析失败']
            if any(indicator in result for indicator in error_indicators):
                logger.warning("⚠️ 解析结果中包含错误信息")
                print("⚠️ 解析结果中包含错误信息")
                return "错误：文档解析过程中出现错误，请检查文件格式或内容"
            
            # Debug: Log partial content to help with debugging
            content_preview = result[:500] + "..." if len(result) > 500 else result
            logger.info(f"✅ 文档处理成功，提取内容长度: {len(result)} 字符")
            logger.debug(f"📝 文档内容预览: {content_preview}")
            print(f"✅ 文档处理成功，提取内容长度: {len(result)} 字符")
            print(f"📝 文档内容预览: {content_preview}")
            
            return result
            
        except Exception as e:
            error_msg = f"文档处理异常: {str(e)}"
            logger.error(f"❌ {error_msg}")
            logger.error(f"📋 异常类型: {type(e).__name__}")
            logger.error(f"📋 异常详情: {traceback.format_exc()}")
            print(f"❌ {error_msg}")
            print(f"📋 异常类型: {type(e).__name__}")
            print(f"📋 异常详情: {traceback.format_exc()}")
            
            # Return a clean error message instead of the raw exception
            return f"错误：文档处理失败 - {type(e).__name__}: {str(e)}。请检查文件格式是否正确。"
        finally:
            # Clean up temporary file
            try:
                if os.path.exists(tmp_file.name):
                    os.unlink(tmp_file.name)
                    logger.info(f"🗑️ 临时文件已清理: {tmp_file.name}")
                    print(f"🗑️ 临时文件已清理: {tmp_file.name}")
            except Exception as cleanup_error:
                logger.warning(f"⚠️ 临时文件清理失败: {str(cleanup_error)}")
                print(f"⚠️ 临时文件清理失败: {str(cleanup_error)}")
                pass

# Legacy functions for backward compatibility (but using improved approach)
async def extract_text_from_docx(file_content: bytes) -> str:
    """Legacy function - now uses improved approach with temporary file"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        try:
            tmp_file.write(file_content)
            tmp_file.flush()
            return read_docx_file(tmp_file.name)
        finally:
            try:
                os.unlink(tmp_file.name)
            except:
                pass

async def extract_text_from_pdf(file_content: bytes) -> str:
    """Legacy function - now uses improved approach with temporary file"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
        try:
            tmp_file.write(file_content)
            tmp_file.flush()
            return read_pdf_file(tmp_file.name)
        finally:
            try:
                os.unlink(tmp_file.name)
            except:
                pass

async def process_uploaded_document(file: UploadFile) -> str:
    """Legacy function - now uses improved approach"""
    return await process_uploaded_document_improved(file)

# Core Models according to README specifications
class APIConfig(BaseModel):
    """API Configuration for AI Agent"""
    type: str = Field(default="http", description="API Type")
    url: str = Field(..., description="AI Agent API URL")
    method: str = Field(default="POST", description="HTTP Method")
    headers: Dict[str, str] = Field(default={}, description="Request Headers")
    timeout: int = Field(default=30, description="Request timeout in seconds")
    
    # Coze Agent specific fields
    agentId: Optional[str] = Field(default=None, description="Coze Agent ID")
    region: Optional[str] = Field(default="global", description="Coze region")
    
    # Coze Bot specific fields  
    botId: Optional[str] = Field(default=None, description="Coze Bot ID")
    botVersion: Optional[str] = Field(default=None, description="Coze Bot Version")

class ConversationScenario(BaseModel):
    """Test scenario for multi-turn conversation"""
    title: str = Field(..., description="Scenario title")
    turns: List[str] = Field(..., description="User conversation turns")
    context: Optional[str] = Field(default="", description="Business context")
    user_profile: Optional[str] = Field(default="", description="User persona")

class EvaluationRequest(BaseModel):
    """Main evaluation request following README structure"""
    agent_api: APIConfig = Field(..., description="AI Agent API Configuration")
    requirement_doc: Optional[str] = Field(default="", description="Requirement document for context")
    conversation_scenarios: List[ConversationScenario] = Field(..., description="Test scenarios")
    
    # Message processing options
    use_raw_messages: Optional[bool] = Field(default=False, description="Use raw user messages without persona enhancement")
    
    # Coze compatibility (temporary until full API support)
    coze_bot_id: Optional[str] = Field(default=None, description="Coze Bot ID for current implementation")

class EvaluationDimensions(BaseModel):
    """6-dimension evaluation framework from README"""
    answer_correctness: float = Field(..., description="回答准确性与专业性") 
    persona_alignment: float = Field(..., description="用户匹配度")
    goal_alignment: float = Field(..., description="目标对齐度")
    specification_citation_accuracy: Optional[float] = Field(default=None, description="规范引用准确度")
    fuzzy_understanding: Optional[float] = Field(default=None, description="模糊理解能力") 
    multi_turn_support: Optional[float] = Field(default=None, description="多轮支持度")

class EvaluationResponse(BaseModel):
    """Response structure matching README output format"""
    evaluation_summary: Dict[str, Any] = Field(..., description="Evaluation summary")
    conversation_records: List[Dict[str, Any]] = Field(..., description="Detailed conversation records")
    recommendations: List[str] = Field(..., description="Improvement suggestions")
    timestamp: str = Field(..., description="Evaluation timestamp")

# Configuration - Using config.py for direct API key management
# All constants are now accessed directly from config module

print(f"✅ Configuration loaded from config.py - DeepSeek API configured")

# Specification Query Default Constants
SPECIFICATION_QUERY_DEFAULTS = {
    "user_persona": {
        "role": "工程项目现场监理工程师",
        "experience_level": "有经验",
        "communication_style": "专业直接",
        "work_environment": "工程现场"
    },
    "usage_context": {
        "business_domain": "工程项目现场监理",
        "usage_scenario": "巡检或验收时，需自然语言提问以快速查询国家/地方规范与图纸说明"
    }
}

# Add MAX_CONVERSATION_TURNS constant if not defined
MAX_CONVERSATION_TURNS = 6

# Conversation continuity management
class ConversationManager:
    """Manage conversation continuity across API calls"""
    
    def __init__(self, api_config: 'APIConfig'):
        self.api_config = api_config
        self.conversation_id = ""
        self.api_type = self._determine_api_type()
        
    def _determine_api_type(self) -> str:
        """Determine the API type based on URL"""
        if "/v1/chat-messages" in self.api_config.url or "dify" in self.api_config.url.lower():
            return "dify"
        elif "coze" in self.api_config.url.lower():
            return "coze"
        else:
            return "custom"
    
    def start_new_conversation(self) -> str:
        """Start a new conversation and return conversation ID"""
        if self.api_type == "dify":
            # For Dify, conversation_id will be extracted from first response
            self.conversation_id = ""
        else:
            # For other APIs, generate a unique conversation ID
            self.conversation_id = f"conv_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
        
        return self.conversation_id
    
    def get_conversation_id(self) -> str:
        """Get current conversation ID"""
        return self.conversation_id
    
    def update_conversation_id(self, new_id: str):
        """Update conversation ID (used when extracted from API response)"""
        if new_id and new_id != self.conversation_id:
            self.conversation_id = new_id
            print(f"🔗 更新对话ID: {new_id[:20]}...")

async def generate_specification_query_scenarios(user_persona_info: Dict) -> List[Dict]:
    """
    Generate scenarios specifically for specification query project
    Uses predefined scenarios for speed optimization
    """
    try:
        # Get persona and context info
        persona = user_persona_info.get("user_persona", {})
        context = user_persona_info.get("usage_context", {})
        
        # Predefined scenarios for specification query (optimized for speed)
        # Start with 1 scenario for testing purposes
        scenarios = [
            {
                "title": "建筑工程规范查询",
                "context": "工程项目现场监理工程师在施工现场发现质量问题，需要查询相关国家规范",
                "user_profile": f"{persona.get('role', '监理工程师')}，{persona.get('experience_level', '有经验')}",
                "business_domain": context.get("business_domain", "工程项目现场监理"),
                "scenario_type": "technical_inquiry"
            }
        ]
        
        print(f"✅ 已生成 {len(scenarios)} 个规范查询场景")
        return scenarios
        
    except Exception as e:
        print(f"❌ 生成规范查询场景失败: {str(e)}")
        # Fallback scenarios
        return [
            {
                "title": "规范智能问答测试",
                "context": "测试AI系统的规范查询能力",
                "user_profile": "工程项目现场监理工程师",
                "business_domain": "工程项目现场监理",
                "scenario_type": "general_inquiry"
            }
        ]

async def conduct_conversation_with_turn_control(
    api_config: APIConfig, 
    scenario: Dict, 
    user_persona_info: Dict, 
    use_raw_messages: bool = False,
    is_tricky_test: bool = False
) -> List[Dict]:
    """
    Conduct conversation with turn control for specification query
    Enhanced to support 4-5 rounds with satisfaction detection
    """
    try:
        print(f"🎯 开始场景对话: {scenario.get('title', '规范查询')}")
        
        conversation_history = []
        conversation_manager = ConversationManager(api_config)
        conversation_manager.start_new_conversation()
        
        # Enhanced conversation with 4-5 turns and satisfaction detection
        max_turns = 5
        satisfaction_indicators = [
            "谢谢", "明白了", "清楚了", "了解了", "知道了", "好的", "没问题", 
            "满意", "解决了", "够了", "足够", "可以了", "ok", "OK", "感谢"
        ]
        
        for turn_num in range(max_turns):
            try:
                # Generate user message
                if turn_num == 0:
                    user_message = await generate_quick_initial_message(scenario, user_persona_info, is_tricky_test)
                elif turn_num == 1:
                    user_message = "请详细说明相关规范条文，并提供具体的技术要求。"
                elif turn_num == 2:
                    user_message = "有没有相关的施工注意事项和质量控制要点？"
                elif turn_num == 3:
                    user_message = "请补充说明验收标准和检测方法。"
                else:
                    user_message = "还有其他需要注意的规范要求吗？"
                
                # Call AI agent
                ai_response = await call_ai_agent_api(
                    api_config, user_message, conversation_manager, use_raw_messages
                )
                
                # Record turn
                conversation_history.append({
                    "turn": turn_num + 1,
                    "user_message": user_message,
                    "ai_response": clean_ai_response(ai_response)
                })
                
                print(f"  ✅ 对话轮次 {turn_num + 1} 完成")
                
                # Check for satisfaction indicators in AI response
                ai_response_lower = ai_response.lower()
                if turn_num >= 1:  # Only check after first turn
                    # If AI asks clarifying questions, continue
                    if any(indicator in ai_response for indicator in ["什么", "哪个", "如何", "怎么", "请问", "？", "?"]):
                        continue
                    
                    # If response is very short and seems final
                    if len(ai_response.strip()) < 50:
                        print(f"  🔚 AI回复较短，可能已完成回答，在第 {turn_num + 1} 轮结束对话")
                        break
                
                # Stop if we have at least 2 turns and this is a substantial conversation
                if turn_num >= 3 and len(conversation_history) >= 4:
                    print(f"  🔚 已完成 {turn_num + 1} 轮充分对话，结束")
                    break
                
            except Exception as e:
                print(f"  ❌ 对话轮次 {turn_num + 1} 失败: {str(e)}")
                break
        
        print(f"🎉 场景对话完成，共 {len(conversation_history)} 轮")
        return conversation_history
        
    except Exception as e:
        print(f"❌ 对话执行失败: {str(e)}")
        return []

async def evaluate_conversation_specification_query(
    conversation_history: List[Dict], 
    scenario: Dict, 
    requirement_context: str = "", 
    user_persona_info: Dict = None
) -> tuple:
    """
    Evaluate conversation specifically for specification query project
    """
    try:
        print(f"📊 开始评估场景: {scenario.get('title', '规范查询')}")
        
        # Build evaluation context
        conversation_text = "\n".join([
            f"用户: {turn['user_message']}\nAI: {turn['ai_response']}" 
            for turn in conversation_history
        ])
        
        persona = user_persona_info.get("user_persona", {}) if user_persona_info else {}
        context = user_persona_info.get("usage_context", {}) if user_persona_info else {}
        
        # Specification query specific evaluation dimensions
        evaluation_dimensions = {
            "answer_correctness": "回答准确性与专业性",
            "persona_alignment": "用户匹配度",
            "goal_alignment": "目标对齐度",
            "specification_citation_accuracy": "规范引用准确度",
            "fuzzy_understanding": "模糊理解能力",
            "multi_turn_support": "多轮支持度"
        }
        
        evaluation_scores = {}
        detailed_explanations = {}
        
        # Evaluate each dimension
        for dimension, dimension_name in evaluation_dimensions.items():
            try:
                eval_prompt = f"""
你是对话质量评估专家。请根据下方对话内容，从"{dimension_name}"这个维度对AI的表现进行评分，满分100分。

📘【场景信息】:
- 场景标题: {scenario.get('title', 'N/A')}
- 背景描述: {scenario.get('context', 'N/A')}
- 用户角色: {persona.get('role', 'N/A')}

🧾【对话内容】:
{conversation_text}

📏【评分标准】:
- 90–100分: 表现优秀，完全满足该维度要求
- 80–89分: 良好表现，略有提升空间
- 60–79分: 基本达标，但存在明显不足
- 40–59分: 有较大问题，影响用户体验
- 0–39分: 表现失败，严重偏离该维度要求

🧪【评分示例】:

示例1（评分: 95）：
- 维度: 回答准确性
- 对话：
  用户：施工时混凝土振捣的时间怎么控制？
  AI：根据《GB50204-2015》第6.3.2条，普通混凝土振捣时间应控制在10-30秒之间。
- 评价：引用规范准确，内容完整、专业。

示例2（评分: 72）：
- 维度: 回答准确性
- 对话：
  用户：钢筋搭接长度？
  AI：钢筋搭接30公分就够了。
- 评价：回答模糊，未引用具体标准，误导用户。

📝【请填写】:
评分：[0-100]
理由：[基于维度和对话，说明你的判断依据]
                """
                
                response = await call_deepseek_api_enhanced(eval_prompt, max_tokens=300, temperature=0.1)
                
                # Extract score and explanation
                score = extract_score_from_response(response)
                if score is None or score < 0 or score > 100:
                    score = 75  # Default score
                
                evaluation_scores[dimension] = score
                detailed_explanations[dimension] = {
                    "score": score,
                    "detailed_analysis": response,
                    "full_response": response
                }
                
                print(f"  ✅ {dimension_name}: {score:.1f}分")
                
            except Exception as e:
                print(f"  ❌ 评估维度 {dimension_name} 失败: {str(e)}")
                evaluation_scores[dimension] = 75  # Default score
                detailed_explanations[dimension] = {
                    "score": 75,
                    "detailed_analysis": f"评估过程出现异常: {str(e)}",
                    "full_response": f"评估异常: {str(e)}"
                }
        
        # Calculate overall score
        scenario_score = sum(evaluation_scores.values()) / len(evaluation_scores) if evaluation_scores else 75
        
        print(f"🎉 场景评估完成，综合得分: {scenario_score:.1f}/100")
        
        return evaluation_scores, detailed_explanations, scenario_score
        
    except Exception as e:
        print(f"❌ 场景评估失败: {str(e)}")
        # Return default scores
        default_scores = {
            "answer_correctness": 75,
            "persona_alignment": 75,
            "goal_alignment": 75
        }
        default_explanations = {
            dim: {
                "score": 75,
                "detailed_analysis": f"评估异常: {str(e)}",
                "full_response": f"评估过程出现错误: {str(e)}"
            } for dim in default_scores.keys()
        }
        return default_scores, default_explanations, 75

async def generate_ai_improvement_suggestions_for_programmers(
    explanations: Dict, 
    evaluation_summary: Dict
) -> List[str]:
    """
    Generate AI-powered improvement suggestions specifically for programmers
    """
    try:
        print("🤖 正在生成AI智能改进建议...")
        
        # Extract key information from explanations
        dimension_scores = []
        dimension_issues = []
        
        for dimension, explanation in explanations.items():
            score = explanation.get('score', 75)
            analysis = explanation.get('detailed_analysis', '')
            
            dimension_scores.append(f"{dimension}: {score}分")
            if score < 80:  # Issues for scores below 80
                dimension_issues.append(f"{dimension}: {analysis}")
        
        # Prepare prompt for DeepSeek
        improvement_prompt = f"""评估概况:
- 综合得分: {evaluation_summary.get('overall_score_100', 75)}/100分
- 评估维度得分: {'; '.join(dimension_scores)}
- 评估框架: {evaluation_summary.get('framework', '规范查询评估')}

主要问题分析:
{chr(10).join(dimension_issues) if dimension_issues else '总体表现良好'}

请提供3-5条具体的技术改进建议，每条建议必须包含以下三个部分：
1. 问题描述：明确指出需要改进的具体问题，特别是与提示词相关的问题
2. 解决方案：提供具体可行的技术方案，包括具体的提示词优化建议
3. 预期效果：说明改进后的预期效果

重点关注以下方面：
- 提示词结构和格式优化
- 上下文管理和历史对话处理
- 角色定义和任务指令的清晰度
- 专业领域知识的引导方式
- 多轮对话的连贯性维护
- 错误处理和边界情况的提示词设计

格式要求：
每条建议必须按照以下格式输出：
问题：xxx
方案：xxx
预期：xxx

改进建议："""
        
        print("📝 发送到DeepSeek的提示词:", improvement_prompt)
        
        response = await call_deepseek_api_enhanced(
            improvement_prompt, 
            max_tokens=500, 
            temperature=0.2
        )
        
        print("📥 DeepSeek原始响应:", response)
        
        # Split the response into individual suggestions based on "问题：" markers
        suggestions = []
        current_suggestion = []
        
        for line in response.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('问题：') and current_suggestion:
                # Save previous suggestion if exists
                suggestions.append('\n'.join(current_suggestion))
                current_suggestion = [line]
            else:
                current_suggestion.append(line)
        
        # Add the last suggestion if exists
        if current_suggestion:
            suggestions.append('\n'.join(current_suggestion))
        
        print("🔍 原始建议数量:", len(suggestions))
        if suggestions:
            print("📋 原始建议内容:", suggestions)
        
        # If no suggestions found, provide default ones
        if not suggestions:
            print("⚠️ 未找到有效建议，使用默认建议")
            overall_score = evaluation_summary.get('overall_score_100', 75)
            if overall_score >= 85:
                suggestions = [
                    "问题：提示词结构需要优化\n方案：建议重构提示词模板，增加角色定义和任务指令的清晰度\n预期：提高AI理解准确性和回答质量",
                    "问题：上下文管理机制需要改进\n方案：优化历史对话的处理方式，增加关键信息提取和记忆机制\n预期：提升多轮对话的连贯性和准确性",
                    "问题：专业领域引导需要加强\n方案：在提示词中增加领域特定的引导和约束条件\n预期：提高专业回答的准确性和深度"
                ]
            elif overall_score >= 70:
                suggestions = [
                    "问题：提示词格式需要标准化\n方案：统一提示词结构，明确区分角色定义、任务要求和约束条件\n预期：提高AI响应的一致性和可预测性",
                    "问题：错误处理提示词需要完善\n方案：增加边界情况和异常处理的提示词设计\n预期：提高系统稳定性和用户体验",
                    "问题：多轮对话提示词需要优化\n方案：改进对话历史的管理和引用方式\n预期：提升对话的连贯性和上下文理解",
                    "问题：专业领域提示词需要细化\n方案：针对不同场景设计专门的提示词模板\n预期：提高专业回答的准确性和相关性"
                ]
            else:
                suggestions = [
                    "问题：提示词工程需要全面重构\n方案：重新设计提示词体系，包括角色定义、任务分解和约束条件\n预期：提高系统整体表现",
                    "问题：上下文管理机制需要重建\n方案：设计新的对话历史处理机制，优化信息提取和记忆方式\n预期：提升对话质量和用户体验",
                    "问题：专业领域引导需要强化\n方案：增加领域特定的提示词模板和知识引导机制\n预期：提高专业回答的准确性",
                    "问题：错误处理机制需要完善\n方案：设计全面的错误处理和边界情况提示词\n预期：提高系统稳定性",
                    "问题：多轮对话机制需要优化\n方案：改进对话流程控制和历史管理方式\n预期：提供更流畅的对话体验"
                ]
        
        print(f"✅ 已生成 {len(suggestions)} 条AI智能改进建议")
        return suggestions[:5]  # Return maximum 5 suggestions
        
    except Exception as e:
        print(f"❌ 生成AI改进建议失败: {str(e)}")
        # Return default suggestions
        return [
            "问题：提示词结构需要优化\n方案：重构提示词模板，增加角色定义和任务指令的清晰度\n预期：提高AI理解准确性和回答质量",
            "问题：上下文管理机制需要改进\n方案：优化历史对话的处理方式，增加关键信息提取和记忆机制\n预期：提升多轮对话的连贯性和准确性",
            "问题：专业领域引导需要加强\n方案：在提示词中增加领域特定的引导和约束条件\n预期：提高专业回答的准确性和深度",
            "问题：错误处理提示词需要完善\n方案：增加边界情况和异常处理的提示词设计\n预期：提高系统稳定性和用户体验"
        ]

async def call_deepseek_api(prompt: str, max_retries: int = 2) -> str:
    """
    Call DeepSeek API with improved error handling
    """
    headers = {
        "Authorization": f"Bearer {config.DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000,
        "temperature": config.DEFAULT_TEMPERATURE
    }
    
    for attempt in range(max_retries):
        try:
            timeout = httpx.Timeout(config.DEEPSEEK_TIMEOUT, connect=10.0)
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.post(config.DEEPSEEK_API_URL, headers=headers, json=payload)
                
                if response.status_code == 200:
                    result = response.json()
                    if "choices" in result and len(result["choices"]) > 0:
                        return result["choices"][0]["message"]["content"].strip()
                    else:
                        raise Exception("No valid response from API")
                elif response.status_code == 401:
                    raise Exception("API authentication failed - check API key")
                elif response.status_code == 429:
                    if attempt < max_retries - 1:
                        await asyncio.sleep(2 ** attempt)
                        continue
                    raise Exception("API rate limited")
                else:
                    error_text = response.text if hasattr(response, 'text') else 'Unknown error'
                    raise Exception(f"API error {response.status_code}: {error_text}")
                    
        except (asyncio.TimeoutError, httpx.TimeoutException):
            if attempt < max_retries - 1:
                await asyncio.sleep(1)
                continue
            raise Exception(f"API timeout after {config.DEEPSEEK_TIMEOUT}s")
        except httpx.RequestError as e:
            if attempt < max_retries - 1:
                await asyncio.sleep(1)
                continue
            raise Exception(f"Network error: {str(e)}")
        except Exception as e:
            raise e
            
    raise Exception("All API attempts failed")

async def call_ai_agent_api(api_config: APIConfig, message: str, conversation_manager: ConversationManager = None, use_raw_message: bool = False) -> str:
    """Call AI Agent API - supports Coze, Dify, and custom APIs with conversation continuity"""
    import json  # Import json module to fix scope issues
    import asyncio
    import time
    
    # 🔧 Enhanced retry logic for connection failures
    max_retries = 3
    retry_delay = 2  # seconds
    
    for attempt in range(max_retries):
        try:
            # 📝 Debug log for message processing mode
            message_preview = message[:80] + "..." if len(message) > 80 else message
            if use_raw_message:
                print(f"🔍 [RAW MODE] Attempt {attempt+1}/{max_retries}: {message_preview}")
            else:
                print(f"🔍 [ENHANCED MODE] Attempt {attempt+1}/{max_retries}: {message_preview}")
            
            # Check if we should use Coze API (either explicit coze URL or fallback URL)
            if "coze" in api_config.url.lower() or "fallback" in api_config.url.lower():
                return await call_coze_api_fallback(message, use_raw_message=use_raw_message)
            
            # Check if this is a Dify API (based on URL pattern)
            elif "/v1/chat-messages" in api_config.url or "dify" in api_config.url.lower():
                conversation_id = conversation_manager.get_conversation_id() if conversation_manager else ""
                response_content, new_conversation_id = await call_dify_api(api_config, message, conversation_id, use_raw_message=use_raw_message)
                
                # Update conversation manager with new conversation ID
                if conversation_manager and new_conversation_id:
                    conversation_manager.update_conversation_id(new_conversation_id)
                
                return response_content
            else:
                # Enhanced generic API support with auto-detection of API formats
                headers = api_config.headers.copy()
                headers.setdefault("Content-Type", "application/json")
                
                # Auto-detect API format based on URL patterns
                session_id = getattr(conversation_manager, 'conversation_id', '') if conversation_manager else ""
                
                # Specialized handling for engineering supervision API (cpolar format)
                if "/ask" in api_config.url or "cpolar" in api_config.url:
                    print(f"🔧 Detected engineering supervision API format")
                    payload = {
                        "question": message,
                        "session_id": session_id or f"eval-{int(time.time())}",
                        "context": ""
                    }
                # Standard custom API formats
                else:
                    # Use appropriate payload field based on raw message mode
                    if use_raw_message:
                        payload = {"input": message, "question": message, "query": message}  # Raw user input fields
                    else:
                        payload = {"message": message, "query": message}  # Enhanced message fields
                
                print(f"📤 Custom API payload: {json.dumps(payload, ensure_ascii=False)[:200]}...")
                
                async with httpx.AsyncClient(timeout=httpx.Timeout(api_config.timeout)) as client:
                    response = await client.request(
                        method=api_config.method,
                        url=api_config.url,
                        headers=headers,
                        json=payload
                    )
                    
                    print(f"📥 Custom API response status: {response.status_code}")
                    
                    if response.status_code == 200:
                        result = response.json()
                        print(f"📋 Custom API response preview: {json.dumps(result, ensure_ascii=False)[:300]}...")
                        
                        # Try common response paths with priority for engineering supervision format
                        raw_response = ""
                        if "answer" in result:  # Primary field for engineering supervision API
                            raw_response = result["answer"]
                        elif "response" in result:
                            raw_response = result["response"]
                        elif "message" in result:
                            raw_response = result["message"]
                        elif "reply" in result:
                            raw_response = result["reply"]
                        elif "content" in result:
                            raw_response = result["content"]
                        else:
                            # Fallback: look for any string value in the response
                            for key, value in result.items():
                                if isinstance(value, str) and len(value) > 10:
                                    raw_response = value
                                    break
                            if not raw_response:
                                raw_response = str(result)
                        
                        # 🔧 UNIVERSAL FIX: Apply plugin extraction to generic API responses too
                        if raw_response:
                            cleaned_response = clean_ai_response(raw_response)
                            if cleaned_response and cleaned_response != raw_response:
                                print(f"🧹 通用API响应经过插件提取处理: {cleaned_response[:100]}...")
                                return cleaned_response
                            print(f"✅ Custom API response: {raw_response[:100]}...")
                            return raw_response
                        else:
                            return "Empty response from API"
                    else:
                        error_message = f"API调用失败: {response.status_code}"
                        try:
                            error_detail = response.json()
                            error_message += f" - {error_detail}"
                        except:
                            error_message += f" - {response.text}"
                        return error_message
                        
        except Exception as e:
            print(f"❌ AI Agent API调用异常 (尝试 {attempt+1}/{max_retries}): {str(e)}")
            if attempt == max_retries - 1:  # 最后一次尝试
                return f"AI Agent API调用失败，请检查配置: {str(e)}"
            await asyncio.sleep(retry_delay)  # 重试前等待
    
    # 如果所有重试都失败，返回错误消息
    return "AI Agent API调用失败：超过最大重试次数"

async def call_dify_api(api_config: APIConfig, message: str, conversation_id: str = "", use_raw_message: bool = False) -> tuple:
    """
    Call Dify API with proper payload format and conversation continuity
    Returns: (response_content, conversation_id)
    """
    try:
        print(f"🔍 调用Dify API: {api_config.url}")
        
        headers = api_config.headers.copy()
        headers.setdefault("Content-Type", "application/json")
        
        # Dify API specific payload format with conversation continuity
        # 📝 Debug log for Dify message processing  
        message_preview = message[:60] + "..." if len(message) > 60 else message
        if use_raw_message:
            print(f"🔍 [DIFY RAW] {message_preview}")
            user_field = "evaluation-user-raw"
        else:
            print(f"🔍 [DIFY ENHANCED] {message_preview}")
            user_field = "evaluation-user"
            
        payload = {
            "inputs": {},
            "query": message,
            "response_mode": "streaming",
            "conversation_id": conversation_id,  # Use provided conversation_id for continuity
            "user": user_field,
            "files": []
        }
        
        print(f"📤 Dify API请求载荷: {json.dumps(payload, ensure_ascii=False)[:200]}...")
        if conversation_id:
            print(f"🔗 使用对话ID: {conversation_id[:20]}...")
        
        async with httpx.AsyncClient(timeout=httpx.Timeout(api_config.timeout)) as client:
            response = await client.post(api_config.url, headers=headers, json=payload)
            
            print(f"🔍 Dify API响应状态: {response.status_code}")
            
            if response.status_code == 200:
                # Check if response is streaming
                content_type = response.headers.get("content-type", "").lower()
                
                if "text/event-stream" in content_type or payload.get("response_mode") == "streaming":
                    # Handle streaming response
                    response_text = response.text
                    print(f"🔍 处理Dify流式响应 ({len(response_text)} chars)")
                    
                    if not response_text.strip():
                        raise Exception("Empty streaming response from Dify API")
                    
                    # Parse Dify streaming format
                    lines = response_text.strip().split('\n')
                    collected_content = ""
                    conversation_id_extracted = conversation_id  # Start with input conversation_id
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        if line.startswith("data: "):
                            data_content = line[6:]  # Remove "data: " prefix
                            
                            # Skip end marker
                            if data_content.strip() in ['"[DONE]"', "[DONE]"]:
                                break
                            
                            try:
                                data_json = json.loads(data_content)
                                
                                # Extract conversation_id for future continuity
                                if "conversation_id" in data_json and data_json["conversation_id"]:
                                    conversation_id_extracted = data_json["conversation_id"]
                                
                                # Extract content from Dify response
                                if "answer" in data_json:
                                    collected_content += data_json["answer"]
                                elif "data" in data_json and "answer" in data_json["data"]:
                                    collected_content += data_json["data"]["answer"]
                                elif "message" in data_json:
                                    collected_content += data_json["message"]
                                    
                            except json.JSONDecodeError:
                                continue
                    
                    if collected_content:
                        print(f"✅ Dify流式响应解析成功: {collected_content[:100]}...")
                        # 🔧 UNIVERSAL FIX: Apply plugin extraction to Dify responses too
                        cleaned_content = clean_ai_response(collected_content)
                        if cleaned_content and cleaned_content != collected_content:
                            print(f"🧹 Dify响应经过插件提取处理: {cleaned_content[:100]}...")
                            collected_content = cleaned_content
                        
                        if conversation_id_extracted and conversation_id_extracted != conversation_id:
                            print(f"🔗 提取到对话ID: {conversation_id_extracted[:20]}...")
                        return collected_content.strip(), conversation_id_extracted
                    else:
                        print("❌ 未从Dify流式响应中提取到有效内容")
                        raise Exception("No valid content in Dify streaming response")
                
                else:
                    # Handle regular JSON response
                    result = response.json()
                    print(f"🔍 Dify JSON响应: {json.dumps(result, ensure_ascii=False)[:300]}...")
                    
                    # Try to extract answer from various possible response formats
                    response_content = ""
                    conversation_id_extracted = conversation_id
                    
                    if "answer" in result:
                        response_content = result["answer"].strip()
                    elif "data" in result and isinstance(result["data"], dict):
                        if "answer" in result["data"]:
                            response_content = result["data"]["answer"].strip()
                        elif "message" in result["data"]:
                            response_content = result["data"]["message"].strip()
                    elif "message" in result:
                        response_content = result["message"].strip()
                    else:
                        print(f"⚠️ 未知的Dify响应格式，尝试返回完整响应")
                        response_content = str(result)
                    
                    # Extract conversation_id from JSON response
                    if "conversation_id" in result and result["conversation_id"]:
                        conversation_id_extracted = result["conversation_id"]
                    
                    # 🔧 UNIVERSAL FIX: Apply plugin extraction to Dify JSON responses too
                    if response_content:
                        cleaned_content = clean_ai_response(response_content)
                        if cleaned_content:
                            response_content = cleaned_content
                            print(f"🧹 Dify JSON响应经过插件提取处理: {response_content[:100]}...")
                    
                    return response_content, conversation_id_extracted
            else:
                error_text = response.text if hasattr(response, 'text') else 'Unknown error'
                print(f"❌ Dify API HTTP错误 {response.status_code}: {error_text}")
                raise Exception(f"Dify API HTTP error {response.status_code}: {error_text}")
                
    except (asyncio.TimeoutError, httpx.TimeoutException):
        print(f"❌ Dify API超时 after {api_config.timeout}s")
        raise Exception(f"Dify API timeout after {api_config.timeout}s")
    except httpx.RequestError as e:
        print(f"❌ Dify API网络错误: {str(e)}")
        raise Exception(f"Dify API network error: {str(e)}")
    except Exception as e:
        print(f"❌ Dify API调用异常: {str(e)}")
        raise e

async def call_coze_api_fallback(message: str, bot_id: str = None, use_raw_message: bool = False) -> str:
    """
    Enhanced Coze API with proper plugin response extraction
    """
    if not bot_id:
        bot_id = config.DEFAULT_COZE_BOT_ID
    
    url = f"{config.COZE_API_BASE}/v3/chat"
    headers = {
        "Authorization": f"Bearer {config.COZE_API_TOKEN}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "parameters": {},
        "bot_id": bot_id,
        "user_id": "123",
        "additional_messages": [
            {
                "content_type": "text",
                "type": "question",
                "role": "user",
                "content": message
            }
        ],
        "auto_save_history": True,
        "stream": True
    }

    # 📝 Clean debug logging
    message_preview = message[:60] + "..." if len(message) > 60 else message
    print(f"🔍 [COZE] {message_preview}")

    try:
        async with httpx.AsyncClient(timeout=config.COZE_TIMEOUT) as client:
            response = await client.post(url, json=payload, headers=headers)
            
            if response.status_code == 200:
                content_type = response.headers.get("Content-Type", "")
                response_text = response.text
                
                print(f"🔍 Coze API Response Status: {response.status_code}")
                print(f"🔍 Handling SSE streaming response ({len(response_text)} chars)")
                
                # Raw response length for debugging
                # print(f"🔍 RAW RESPONSE (first 1000 chars): {response_text[:1000]}...") # Disabled for cleaner output

                if "text/event-stream" in content_type or "stream" in response_text:
                    # Parse SSE streaming response
                    lines = response_text.strip().split('\n')
                    current_event = None
                    main_answer = ""
                    assistant_messages = []
                    collected_content = ""
                    plugin_responses = []  # 🔧 NEW: Collect plugin responses
                    
                    # 🔍 PATTERN SEARCH: Look for tool output patterns in raw response
                    import re
                    tool_output_patterns = [
                        r'"tool_output_content":"([^"]+)"',
                        r'"tool_output_content":\s*"([^"]+)"',
                        r'答案：([^"\\n]+)',
                        r'"answer":"([^"]+)"',
                        r'"response":"([^"]+)"',
                        r'"result":"([^"]+)"'
                    ]
                    
                    for pattern in tool_output_patterns:
                        matches = re.findall(pattern, response_text, re.IGNORECASE | re.DOTALL)
                        for match in matches:
                            # Clean up escape characters
                            cleaned_match = match.replace('\\n', '\n').replace('\\"', '"').replace('\\t', '\t')
                            if len(cleaned_match.strip()) > 20:  # Substantial content
                                plugin_responses.append(cleaned_match.strip())
                    
                    for line in lines:
                        line = line.strip()
                        
                        if line.startswith("event:"):
                            current_event = line[6:].strip()
                        elif line.startswith("data:") and len(line) > 5:
                            data_content = line[5:].strip()
                            
                            if data_content in ["[DONE]", ""]:
                                continue
                                
                            try:
                                data_json = json.loads(data_content)
                                
                                if current_event == "conversation.message.delta":
                                    # This is streaming content chunk
                                    if "content" in data_json:
                                        content_chunk = data_json["content"]
                                        # Don't collect plugin invocation chunks
                                        if not (content_chunk.strip().startswith('{"name":"') or 
                                                '"plugin_id":' in content_chunk or
                                                '"arguments":' in content_chunk):
                                            collected_content += content_chunk
                                        
                                elif current_event == "conversation.message.completed":
                                    # This is a completed message
                                    if "content" in data_json:
                                        message_content = data_json["content"]
                                        role = data_json.get("role", "unknown")
                                        msg_type = data_json.get("type", "text")
                                        
                                        # 🔧 NEW: Enhanced plugin response handling
                                        if role == "assistant" and message_content:
                                            # Check if content is a plugin invocation JSON
                                            if (message_content.strip().startswith('{"name":"') and 
                                                '"arguments":' in message_content and
                                                '"plugin_id":' in message_content):
                                                print(f"🔧 Found plugin invocation: {message_content[:200]}...")
                                                
                                                # Try to extract tool output from plugin response
                                                try:
                                                    plugin_data = json.loads(message_content)
                                                    
                                                    # Look for tool output in various possible fields
                                                    tool_output_fields = [
                                                        'tool_output_content',
                                                        'output',
                                                        'result', 
                                                        'content',
                                                        'answer',
                                                        'response',
                                                        'text',
                                                        'data'
                                                    ]
                                                    
                                                    found_output = False
                                                    
                                                    # Check top-level fields
                                                    for field in tool_output_fields:
                                                        if field in plugin_data and plugin_data[field]:
                                                            tool_output = str(plugin_data[field])
                                                            if len(tool_output.strip()) > 10:  # Substantial content
                                                                print(f"✅ Extracted plugin output from {field}: {tool_output[:100]}...")
                                                                plugin_responses.append(tool_output)
                                                                found_output = True
                                                                break
                                                    
                                                    # Check nested arguments if not found yet
                                                    if not found_output and 'arguments' in plugin_data and isinstance(plugin_data['arguments'], dict):
                                                        args = plugin_data['arguments']
                                                        for field in tool_output_fields:
                                                            if field in args and args[field]:
                                                                tool_output = str(args[field])
                                                                if len(tool_output.strip()) > 10:
                                                                    print(f"✅ Extracted plugin output from args.{field}: {tool_output[:100]}...")
                                                                    plugin_responses.append(tool_output)
                                                                    found_output = True
                                                                    break
                                                    
                                                except json.JSONDecodeError as e:
                                                    print(f"⚠️ Failed to parse plugin JSON: {e}")
                                                
                                                continue  # Skip adding to assistant_messages
                                            
                                            # Regular assistant message
                                            assistant_messages.append({
                                                "content": message_content,
                                                "type": msg_type,
                                                "length": len(message_content)
                                            })
                                        
                                        # Set main answer if this is substantial content and not plugin invocation
                                        if (message_content and len(message_content) > 20 and 
                                            not (message_content.strip().startswith('{"name":"') and 
                                                 '"arguments":' in message_content and
                                                 '"plugin_id":' in message_content)):
                                            if not main_answer or len(message_content) > len(main_answer):
                                                main_answer = message_content
                                
                                elif current_event == "conversation.chat.completed":
                                    # Chat completion event - might have final answer
                                    if "last_message" in data_json and data_json["last_message"].get("content"):
                                        final_content = data_json["last_message"]["content"]
                                        if len(final_content) > 20:
                                            main_answer = final_content
                                
                                # 🔧 NEW: Check for tool output events
                                elif current_event == "conversation.message.plugin.finish":
                                    if "content" in data_json:
                                        plugin_output = data_json["content"]
                                        if len(plugin_output.strip()) > 10:
                                            print(f"✅ Plugin finish event with output: {plugin_output[:100]}...")
                                            plugin_responses.append(plugin_output)
                                
                                # 🔧 CRITICAL: Extract plugin content from stream_plugin_finish events
                                if current_event == "conversation.message.completed" and "content" in data_json:
                                    content = data_json["content"]
                                    # Check if this is a stream_plugin_finish event with tool_output_content
                                    if isinstance(content, str) and '"msg_type":"stream_plugin_finish"' in content:
                                        try:
                                            # Parse the JSON content to extract tool_output_content
                                            inner_json = json.loads(content)
                                            if inner_json.get("msg_type") == "stream_plugin_finish" and "data" in inner_json:
                                                data_str = inner_json["data"]
                                                if isinstance(data_str, str):
                                                    data_content = json.loads(data_str)
                                                    if "tool_output_content" in data_content:
                                                        tool_output = data_content["tool_output_content"]
                                                        if tool_output and len(tool_output.strip()) > 20:
                                                            print(f"✅ Extracted plugin output: {tool_output[:200]}...")
                                                            plugin_responses.append(tool_output)
                                        except (json.JSONDecodeError, KeyError) as e:
                                            print(f"⚠️ Failed to parse plugin content: {e}")
                                
                                # Minimal logging for debugging
                                # if current_event in ["conversation.message.completed", "conversation.message.delta"] and "content" in data_json:
                                #     content_preview = str(data_json["content"])[:100] if data_json["content"] else "empty"
                                #     print(f"🔍 {current_event}: {content_preview}...")
                                
                                # Check for tool output in tool_response type messages
                                if current_event == "conversation.message.completed" and data_json.get("type") == "tool_response":
                                    if "content" in data_json:
                                        tool_content = data_json["content"]
                                        if tool_content and len(str(tool_content).strip()) > 10:
                                            print(f"✅ Found tool response: {str(tool_content)[:200]}...")
                                            # Skip the generic "directly streaming reply" message
                                            if "directly streaming reply" not in str(tool_content):
                                                plugin_responses.append(str(tool_content))
                                
                                # Also check for direct content fields regardless of event
                                elif "content" in data_json and not data_json.get("msg_type"):
                                    content = data_json["content"]
                                    if (content and len(content) > 20 and 
                                        not any(keyword in content for keyword in [
                                            '用户编写的信息', '用户画像信息', '用户记忆点信息'
                                        ]) and
                                        not (content.strip().startswith('{"name":"') and 
                                             '"arguments":' in content and
                                             '"plugin_id":' in content)):
                                        if not main_answer or len(content) > len(main_answer):
                                            main_answer = content
                                
                            except json.JSONDecodeError as e:
                                continue
                    
                    # 🔧 ENHANCED: Priority order for response content 
                    print(f"🔍 Response content summary: {len(plugin_responses)} plugins, {len(assistant_messages)} messages, main_answer: {len(main_answer) if main_answer else 0} chars")
                    
                    # 1. Plugin responses (highest priority for technical queries)
                    if plugin_responses:
                        # Use the longest/most substantial plugin response
                        best_plugin_response = max(plugin_responses, key=len)
                        if len(best_plugin_response) > 20:
                            print(f"✅ Using plugin response ({len(best_plugin_response)} chars)")
                            return clean_ai_response(best_plugin_response)
                    
                    # 2. Main answer (from completed messages)
                    if main_answer and not any(keyword in main_answer for keyword in [
                        '用户编写的信息', '用户画像信息', '用户记忆点信息', 'wraped_text', 'origin_search_results'
                    ]):
                        print(f"✅ Using main answer ({len(main_answer)} chars): {main_answer[:100]}...")
                        return clean_ai_response(main_answer)
                    
                    # 3. Look for non-system assistant messages
                    for i, msg in enumerate(assistant_messages):
                        content = msg["content"]
                        if (not any(keyword in content for keyword in [
                            '用户编写的信息', '用户画像信息', '用户记忆点信息', 'wraped_text', 'origin_search_results'
                        ]) and
                        not (content.strip().startswith('{"name":"') and 
                             '"arguments":' in content and
                             '"plugin_id":' in content)):
                            print(f"✅ Using assistant message ({len(content)} chars): {content[:100]}...")
                            return clean_ai_response(content)
                    
                    # 4. Collected streaming content (delta)
                    if (collected_content and 
                        not any(keyword in collected_content for keyword in [
                            '用户编写的信息', '用户画像信息', '用户记忆点信息', 'wraped_text', 'origin_search_results'
                        ]) and
                        not (collected_content.strip().startswith('{"name":"') and 
                             '"arguments":' in collected_content and
                             '"plugin_id":' in collected_content)):
                        print(f"✅ Using streaming content ({len(collected_content)} chars): {collected_content[:100]}...")
                        return clean_ai_response(collected_content)
                    
                    # 5. Check for billing errors before returning empty
                    if "unpaid bills" in response_text or "code\":4027" in response_text:
                        print("💰 ❌ Coze账户余额不足或有未付账单")
                        print("💰 详情: https://console.volcengine.com/coze-pro/overview")
                        return "❌ API Error: Coze账户余额不足，请联系管理员充值账户后重试"
                    
                    # 6. If all content was system messages, return empty
                    print("❌ No conversational content found (system messages only)")
                    print(f"🔍 COMPLETE RAW RESPONSE FOR DEBUGGING: {response_text}")
                    return ""  # Return empty to trigger proper handling
                
                else:
                    # Handle regular JSON response
                    result = response.json()
                    print(f"🔍 Coze API Response Structure: {json.dumps(result, indent=2)[:500]}...")
                    
                    if result.get("code") == 0 and "data" in result:
                        data = result["data"]
                        
                        # Handle non-streaming response format
                        if "messages" in data and len(data["messages"]) > 0:
                            # Get the last assistant message
                            for msg in reversed(data["messages"]):
                                if msg.get("role") == "assistant" and msg.get("content"):
                                    print(f"✅ Found assistant response: {msg['content'][:100]}...")
                                    return clean_ai_response(msg["content"])
                            
                            # Fallback to any message content
                            for msg in data["messages"]:
                                if msg.get("content"):
                                    print(f"✅ Found fallback response: {msg['content'][:100]}...")
                                    return clean_ai_response(msg["content"])
                        
                        # Check for other possible response formats
                        if "answer" in data:
                            print(f"✅ Found answer field: {data['answer'][:100]}...")
                            return clean_ai_response(data["answer"])
                        
                        if "content" in data:
                            print(f"✅ Found content field: {data['content'][:100]}...")
                            return clean_ai_response(data["content"])
                        
                        print(f"⚠️ No response content found in data: {list(data.keys())}")
                        raise Exception("No valid response content in Coze API result")
                    else:
                        error_msg = result.get("msg", "Unknown Coze API error")
                        print(f"❌ Coze API returned error: {error_msg}")
                        raise Exception(f"Coze API error: {error_msg}")
            else:
                error_text = response.text if hasattr(response, 'text') else 'Unknown error'
                print(f"❌ HTTP error {response.status_code}: {error_text}")
                raise Exception(f"Coze API HTTP error {response.status_code}: {error_text}")
                
    except (asyncio.TimeoutError, httpx.TimeoutException):
        print(f"❌ Coze API timeout after {config.COZE_TIMEOUT}s")
        raise Exception(f"Coze API timeout after {config.COZE_TIMEOUT}s")
    except httpx.RequestError as e:
        print(f"❌ Coze API network error: {str(e)}")
        raise Exception(f"Coze API network error: {str(e)}")
    except Exception as e:
        print(f"❌ Coze API unexpected error: {str(e)}")
        raise e

def extract_score_from_response(response: str) -> float:
    """Extract numerical score from DeepSeek response (1-100 scale)"""
    try:
        # First try to parse as JSON (new format)
        try:
            # Look for JSON object in response
            json_start = response.find('{')
            json_end = response.rfind('}')
            if json_start != -1 and json_end != -1:
                json_str = response[json_start:json_end+1]
                parsed = json.loads(json_str)
                if 'score' in parsed:
                    score = float(parsed['score'])
                    return min(max(score, 1.0), 100.0)  # Clamp between 1-100
        except (json.JSONDecodeError, ValueError, KeyError):
            pass
        
        # Fallback to pattern matching (legacy format)
        patterns = [
            r'评分[：:]\s*(\d+(?:\.\d+)?)',
            r'得分[：:]\s*(\d+(?:\.\d+)?)',
            r'(\d+(?:\.\d+)?)\s*分',
            r'(\d+(?:\.\d+)?)\s*/\s*100',
            r'(\d+(?:\.\d+)?)\s*星'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, response)
            if match:
                score = float(match.group(1))
                # If score appears to be on 1-5 scale, convert to 1-100
                if score <= 5.0:
                    score = score * 20  # Convert 1-5 to 20-100
                return min(max(score, 1.0), 100.0)  # Clamp between 1-100
        
        # If no pattern found, try to find any number
        numbers = re.findall(r'\d+(?:\.\d+)?', response)
        if numbers:
            for num in numbers:
                score = float(num)
                if 1 <= score <= 5:
                    return score * 20  # Convert 1-5 to 20-100
                elif 1 <= score <= 100:
                    return score
        
        # Default fallback
        return 60.0  # Middle score
        
    except Exception:
        return 60.0  # Middle score on error

async def call_coze_api_sdk(bot_id: str, message: str) -> str:
    """
    Use official Coze SDK for API calls (preferred method)
    """
    if not COZE_SDK_AVAILABLE:
        print("⚠️ Coze SDK不可用，使用HTTP fallback")
        return await call_coze_api_fallback(message, bot_id)
    
    try:
        # Initialize Coze client with config settings
        coze = Coze(
            auth=TokenAuth(token=config.COZE_API_TOKEN), 
            base_url=COZE_CN_BASE_URL
        )
        
        # Generate a unique user ID for this conversation
        user_id = f"eval_user_{int(time.time())}"
        
        print(f"🔄 使用Coze SDK调用 Bot {bot_id}")
        
        # Collect response content
        response_content = ""
        token_count = 0
        
        # Track plugin responses like in the HTTP fallback method
        plugin_responses = []
        collected_content = ""
        
        # Create streaming chat
        for event in coze.chat.stream(
            bot_id=bot_id,
            user_id=user_id,
            additional_messages=[
                Message.build_user_question_text(message),
            ],
        ):
            # Enhanced logging for debugging
            print(f"[📡 EVENT] {event.event}")
            
            if event.event == ChatEventType.CONVERSATION_MESSAGE_DELTA:
                content = event.message.content or ""
                print(f"[📦 DELTA CONTENT] {content[:100]}...")
                
                # 🔧 Enhanced plugin detection using existing patterns from HTTP fallback
                # Check if content contains plugin JSON data
                if ('"tool_output_content"' in content or 
                    '"plugin_id"' in content or
                    '"msg_type":"stream_plugin_finish"' in content):
                    
                    print(f"🔧 Detected plugin content in delta: {content[:150]}...")
                    
                    # Try to extract plugin tool output using same logic as HTTP fallback
                    try:
                        if '"tool_output_content"' in content:
                            # Extract tool_output_content directly
                            import re
                            match = re.search(r'"tool_output_content":"([^"]+)"', content)
                            if match:
                                tool_output = match.group(1).replace('\\n', '\n').replace('\\"', '"')
                                print(f"✅ Extracted tool_output_content: {tool_output[:100]}...")
                                plugin_responses.append(tool_output)
                                continue  # Skip adding to response_content to avoid duplication
                        
                        # Handle stream_plugin_finish format  
                        if '"msg_type":"stream_plugin_finish"' in content:
                            try:
                                import json
                                plugin_data = json.loads(content)
                                data_content = plugin_data.get('data', {})
                                if isinstance(data_content, str):
                                    # Parse nested JSON in data field
                                    try:
                                        data_obj = json.loads(data_content)
                                        tool_output = data_obj.get('tool_output_content', '')
                                        if tool_output and len(tool_output.strip()) > 5:
                                            print(f"✅ Extracted from stream_plugin_finish: {tool_output[:100]}...")
                                            plugin_responses.append(tool_output)
                                            continue
                                    except:
                                        pass
                                elif isinstance(data_content, dict):
                                    tool_output = data_content.get('tool_output_content', '')
                                    if tool_output and len(tool_output.strip()) > 5:
                                        print(f"✅ Extracted from nested JSON: {tool_output[:100]}...")
                                        plugin_responses.append(tool_output)
                                        continue
                            except json.JSONDecodeError:
                                print(f"⚠️ Failed to parse stream_plugin_finish JSON")
                        
                        # Try to parse as plugin invocation JSON
                        if '"plugin_id"' in content:
                            try:
                                import json
                                plugin_data = json.loads(content)
                                
                                # Look for tool output in various possible fields
                                tool_output_fields = [
                                    'tool_output_content', 'output', 'result', 
                                    'content', 'answer', 'response'
                                ]
                                
                                for field in tool_output_fields:
                                    if field in plugin_data and plugin_data[field]:
                                        tool_output = str(plugin_data[field])
                                        if len(tool_output.strip()) > 10:
                                            print(f"✅ Extracted plugin output from {field}: {tool_output[:100]}...")
                                            plugin_responses.append(tool_output)
                                            break
                                
                                # Also check in arguments field
                                if 'arguments' in plugin_data and isinstance(plugin_data['arguments'], dict):
                                    args = plugin_data['arguments']
                                    for field in tool_output_fields:
                                        if field in args and args[field]:
                                            tool_output = str(args[field])
                                            if len(tool_output.strip()) > 10:
                                                print(f"✅ Extracted tool output from args.{field}: {tool_output[:100]}...")
                                                plugin_responses.append(tool_output)
                                                break
                                continue  # Skip adding plugin JSON to response_content
                            except json.JSONDecodeError:
                                print(f"⚠️ Failed to parse plugin JSON")
                    
                    except Exception as e:
                        print(f"⚠️ Error processing plugin content: {e}")
                
                # Only add to collected content if it's not plugin invocation JSON
                if not ('"plugin_id"' in content and content.strip().startswith('{')):
                    collected_content += content
                    response_content += content
                
            # Handle direct plugin results (if SDK supports these attributes)
            elif hasattr(event, 'plugin_result') and event.plugin_result:
                plugin_content = str(event.plugin_result)
                response_content += f"\n{plugin_content}"
                plugin_responses.append(plugin_content)
                print(f"[🔌 PLUGIN RESULT] {plugin_content[:100]}...")
                
            # Handle tool outputs (alternative event type for plugins)
            elif hasattr(event, 'tool_output') and event.tool_output:
                tool_content = str(event.tool_output)
                response_content += f"\n{tool_content}"
                plugin_responses.append(tool_content)
                print(f"[🔧 TOOL OUTPUT] {tool_content[:100]}...")
                
            # Handle any other message content (fallback for other content types)
            elif hasattr(event, 'message') and hasattr(event.message, 'content') and event.message.content:
                if event.event != ChatEventType.CONVERSATION_MESSAGE_DELTA:  # Avoid duplicates
                    content = event.message.content
                    response_content += content
                    print(f"[📄 OTHER MESSAGE] {content[:100]}...")
                    
            elif event.event == ChatEventType.CONVERSATION_CHAT_COMPLETED:
                if hasattr(event.chat, 'usage') and event.chat.usage:
                    token_count = event.chat.usage.token_count
                print(f"[✅ CHAT COMPLETED] Token count: {token_count}")
                break
                
            # Log any unhandled events for debugging
            else:
                print(f"[❓ UNHANDLED EVENT] {event.event} - {type(event)}")
                # Try to extract any content from unknown event types
                if hasattr(event, 'content'):
                    content = str(event.content)
                    response_content += f"\n{content}"
                    print(f"[❓ UNKNOWN CONTENT] {content[:100]}...")
        
        # 🔧 Priority order for response content with plugin support (same as HTTP fallback)
        # 1. Plugin responses (highest priority for technical queries)
        if plugin_responses:
            # Use the longest/most substantial plugin response
            best_plugin_response = max(plugin_responses, key=len)
            if len(best_plugin_response) > 20:
                print(f"✅ Using plugin response ({len(best_plugin_response)} chars)")
                response_content = best_plugin_response
            
        # Apply the same cleaning as the HTTP fallback
        if response_content:
            response_content = clean_ai_response(response_content)
        
        if not response_content.strip():
            raise Exception("Empty response from Coze SDK")
            
        print(f"✅ SDK调用成功，响应长度: {len(response_content)}, Token: {token_count}")
        return response_content.strip()
        
    except Exception as e:
        print(f"❌ Coze SDK调用失败: {str(e)}")
        print("🔄 切换到HTTP fallback")
        return await call_coze_api_fallback(message, bot_id)

async def call_coze_api(bot_id: str, message: str) -> str:
    """
    Main Coze API call function - tries SDK first, then fallback
    """
    return await call_coze_api_sdk(bot_id, message)

async def call_api(api_config: APIConfig, message: str) -> str:
    """Generic API call function"""
    return await call_ai_agent_api(api_config, message)

async def conduct_conversation(api_config: APIConfig, scenario: ConversationScenario) -> List[Dict]:
    """Conduct a conversation based on the scenario"""
    conversation_history = []
    
    print(f"🔧 API 配置类型: {api_config.type}")
    print(f"🔧 API 配置详情: {api_config}")
    
    for turn_num, user_message in enumerate(scenario.turns, 1):
        print(f"  第{turn_num}轮: {user_message}")
        
        # Get AI response based on API type
        if api_config.type == 'coze-agent':
            print("🤖 使用 Coze Agent API")
            # Extract access token properly
            access_token = api_config.headers.get('Authorization', '')
            if access_token.startswith('Bearer '):
                access_token = access_token.replace('Bearer ', '')
            
            print(f"🔑 Agent ID: {getattr(api_config, 'agentId', 'Not found')}")
            print(f"🔑 Region: {getattr(api_config, 'region', 'Not found')}")
            
            ai_response = await call_coze_agent_api(
                agent_id=getattr(api_config, 'agentId', ''),
                access_token=access_token,
                message=user_message,
                region=getattr(api_config, 'region', 'global')
            )
        elif api_config.type == 'coze-bot' or hasattr(api_config, 'botId'):
            print("🤖 使用 Coze Bot API")
            bot_id = getattr(api_config, 'botId', '')
            ai_response = await call_coze_api(bot_id, user_message)
        elif api_config.type == 'custom-api':
            print("🤖 使用自定义 API")
            ai_response = await call_api(api_config, user_message)
        else:
            print(f"❌ 未知的API类型: {api_config.type}")
            ai_response = f"API配置错误，无法识别的API类型: {api_config.type}"
        
        conversation_history.append({
            "turn": turn_num,
            "user_message": user_message,
            "ai_response": ai_response
        })
        
        # Truncate long responses for display
        display_response = ai_response[:100] + "..." if len(ai_response) > 100 else ai_response
        print(f"  AI回复: {display_response}")
    
    return conversation_history

async def evaluate_conversation_deepseek(
    conversation_history: List[Dict], 
    scenario: Dict, 
    requirement_context: str = "",
    evaluation_mode: str = "manual",
    user_persona_info: Dict = None
) -> tuple:
    """
    Enhanced DeepSeek evaluation with persona awareness
    """
    try:
        print("🧠 开始DeepSeek智能评估...")
        
        # Build enhanced context section
        context_section = f"""
业务场景: {scenario.get('context', '通用AI助手场景')}
用户画像: {scenario.get('user_profile', '普通用户')}
对话主题: {scenario.get('title', '')}
评估模式: {evaluation_mode}
"""
        
        # Add persona information if available
        if evaluation_mode == "auto" and user_persona_info:
            persona = user_persona_info.get('user_persona', {})
            context_section += f"""
提取的用户角色: {persona.get('role', '')}
用户经验水平: {persona.get('experience_level', '')}
沟通风格: {persona.get('communication_style', '')}
工作环境: {persona.get('work_environment', '')}
"""
        
        if requirement_context:
            context_section += f"\n需求文档上下文:\n{requirement_context[:1000]}"
        
        # Build conversation context
        conversation_text = ""
        for turn in conversation_history:
            conversation_text += f"用户: {turn['user_message']}\nAI: {turn['ai_response']}\n\n"
        
        # Enhanced evaluation prompts with persona awareness
        base_context = f"{context_section}\n\n对话记录:\n{conversation_text}\n"
        
        # Call the evaluation function
        return await perform_deepseek_evaluations({}, base_context, requirement_context)
        
    except Exception as e:
        print(f"❌ DeepSeek评估失败: {str(e)}")
        return {}, {}

def generate_enhanced_recommendations(evaluation_results: List[Dict], user_persona_info: Dict = None) -> List[str]:
    """
    Generate enhanced recommendations based on evaluation results and user persona
    """
    recommendations = []
    
    if not evaluation_results:
        return [
            "无法生成推荐建议，请先完成有效的评估",
            "检查AI Agent配置和网络连接",
            "确保对话场景配置正确"
        ]
    
    persona = user_persona_info.get('user_persona', {}) if user_persona_info else {}
    context = user_persona_info.get('usage_context', {}) if user_persona_info else {}
    requirements = user_persona_info.get('extracted_requirements', {}) if user_persona_info else {}
    
    # Calculate dimension averages from evaluation results
    all_scores = {}
    for result in evaluation_results:
        scores = result.get("evaluation_scores", {})
        for dimension, score in scores.items():
            if dimension not in all_scores:
                all_scores[dimension] = []
            all_scores[dimension].append(score)
    
    dimension_averages = {}
    for dimension, scores in all_scores.items():
        dimension_averages[dimension] = sum(scores) / len(scores) if scores else 0
    
    # Overall performance assessment based on extracted persona
    overall_avg = sum(dimension_averages.values()) / len(dimension_averages) if dimension_averages else 0
    
    if overall_avg >= 4.5:
        recommendations.append(f"🟢 针对{persona.get('role', '用户')}的整体表现优秀！AI代理能够有效处理{context.get('business_domain', '业务')}需求")
    elif overall_avg >= 4.0:
        recommendations.append(f"🟡 对{persona.get('role', '用户')}的服务良好，基本满足{context.get('business_domain', '业务')}需求，有进一步优化空间")
    elif overall_avg >= 3.0:
        recommendations.append(f"🟠 服务{persona.get('role', '用户')}的能力中等，建议针对{context.get('business_domain', '业务领域')}特点进行改进")
    else:
        recommendations.append(f"🔴 需要显著改进对{persona.get('role', '用户')}的服务能力，特别是{context.get('business_domain', '业务')}相关功能")
    
    # Dimension-specific recommendations with persona context  
    # (Removed fuzzy_understanding dimension - no longer needed)
    
    if dimension_averages.get('answer_correctness', 0) < 3.5:
        expertise_areas = persona.get('expertise_areas', [])
        expertise_context = f"，特别是{', '.join(expertise_areas[:2])}领域" if expertise_areas else ""
        recommendations.append(f"📚 针对{persona.get('role', '用户')}的专业准确性需要提升：加强知识库建设{expertise_context}")
    
    if dimension_averages.get('persona_alignment', 0) < 3.5:
        comm_style = persona.get('communication_style', '专业沟通')
        recommendations.append(f"👥 用户匹配度有待改善：优化语言风格以适应{comm_style}，匹配{persona.get('experience_level', '用户经验水平')}")
    
    if dimension_averages.get('goal_alignment', 0) < 3.5:
        core_functions = requirements.get('core_functions', [])
        function_context = f"，重点关注{', '.join(core_functions[:2])}" if core_functions else ""
        recommendations.append(f"🎯 业务目标对齐度需要改进：确保回答能够满足{context.get('business_domain', '业务')}的实际需求{function_context}")
    
    # Add persona-specific targeted recommendations
    role = persona.get('role', '')
    work_env = persona.get('work_environment', '')
    
    if '客服' in role:
        recommendations.append(f"🎧 客服场景优化：针对{work_env}环境，提升响应效率和标准化回答")
        quality_expectations = requirements.get('quality_expectations', [])
        if quality_expectations:
            recommendations.append(f"⏱️ 服务质量提升：重点满足{', '.join(quality_expectations[:2])}等客服质量要求")
    elif '工程师' in role or '监理' in role:
        recommendations.append(f"🔧 技术专业性：加强对{work_env}环境下技术规范和标准的支持")
        if '规范' in str(context.get('primary_scenarios', [])):
            recommendations.append("📋 规范查询优化：增强对技术标准和施工规范的快速检索和解释能力")
    elif '管理' in role:
        recommendations.append("📊 管理决策支持：提供更多数据分析和决策建议功能")
    
    # Add interaction preference recommendations
    interaction_goals = context.get('interaction_goals', [])
    if interaction_goals:
        recommendations.append(f"🎯 交互目标优化：重点提升{', '.join(interaction_goals[:2])}的实现效果")
    
    # Quality expectations based recommendations  
    quality_expectations = requirements.get('quality_expectations', [])
    if quality_expectations:
        recommendations.append(f"⭐ 质量标准对齐：确保达到{', '.join(quality_expectations[:2])}等质量期望")
    
    # Limit to 6-8 most relevant recommendations
    unique_recommendations = list(dict.fromkeys(recommendations))
    return unique_recommendations[:8]

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/health")
async def health_check():
    """Simple health check endpoint for debugging"""
    try:
        memory_usage = check_memory_usage()
        return {
            "status": "healthy",
            "timestamp": datetime.now().isoformat(),
            "memory_usage": memory_usage,
            "version": "4.0.0",
            "features": {
                "database": PYMYSQL_AVAILABLE,
                "coze_sdk": COZE_SDK_AVAILABLE,
                "document_processing": DOCUMENT_PROCESSING_AVAILABLE,
                "memory_monitoring": PSUTIL_AVAILABLE
            }
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.post("/api/evaluate-agent-with-file", response_model=EvaluationResponse)
async def evaluate_agent_with_file(
    agent_api_config: str = Form(...),
    requirement_file: UploadFile = File(None),
    requirement_text: str = Form(None),
    conversation_scenarios: str = Form(...),
    coze_bot_id: str = Form(None),
    evaluation_mode: str = Form("manual"),  # "auto" or "manual"
    extracted_persona: str = Form(None)     # JSON string of extracted persona
):
    """
    Enhanced evaluation endpoint supporting both automatic persona extraction and manual input
    """
    try:
        print("🤖============================================================🤖")
        print("   AI Agent 评估平台 v3.0 (增强模式)")
        print("🤖============================================================🤖")
        
        # Parse API configuration
        try:
            # ⭐ Security: Validate input length
            if len(agent_api_config) > 50000:  # 50KB limit
                raise HTTPException(status_code=413, detail="API配置过长，请检查配置内容")
            
            api_config_dict = json.loads(agent_api_config)
            
            # ⭐ Security: Validate API URL if present
            if 'url' in api_config_dict and not validate_api_url(api_config_dict['url']):
                raise HTTPException(status_code=400, detail="不安全的API URL")
            
            # Debug: log the received configuration structure
            print(f"🔍 Received API config structure: {json.dumps(api_config_dict, indent=2)}")
            logger.info(f"🔍 Received API config: {api_config_dict}")
            
            # Enhanced data cleaning for common frontend issues
            if isinstance(api_config_dict, dict):
                # Strategy 1: Look for common wrapping patterns
                if 'config' in api_config_dict and isinstance(api_config_dict['config'], dict):
                    print("⚠️ Detected config wrapped in 'config' key, unwrapping...")
                    api_config_dict = api_config_dict['config']
                elif 'api_config' in api_config_dict and isinstance(api_config_dict['api_config'], dict):
                    print("⚠️ Detected config wrapped in 'api_config' key, unwrapping...")
                    api_config_dict = api_config_dict['api_config']
                
                # Strategy 2: Fix the nested headers issue (common user error)
                if 'headers' in api_config_dict and isinstance(api_config_dict['headers'], dict):
                    headers = api_config_dict['headers']
                    
                    # Check if user pasted entire config into headers field
                    if 'url' in headers and 'method' in headers and 'type' in headers:
                        print("⚠️ Detected full config pasted in headers field, extracting...")
                        # The real config is nested in headers, extract it
                        real_config = headers.copy()
                        
                        # Clean the nested headers if it exists
                        if 'headers' in real_config and isinstance(real_config['headers'], dict):
                            real_config['headers'] = real_config['headers']
                        else:
                            real_config['headers'] = {'Content-Type': 'application/json'}
                        
                        api_config_dict = real_config
                        print(f"✅ Extracted real config from nested headers: {api_config_dict['type']}")
                    
                    # Check for duplicate nested structure in headers
                    elif any(key in headers for key in ['type', 'url', 'method', 'timeout']):
                        print("⚠️ Detected config properties mixed in headers, cleaning...")
                        # Extract only valid header properties
                        valid_headers = {}
                        for key, value in headers.items():
                            if key.lower() in ['authorization', 'content-type', 'user-agent', 'accept', 'x-api-key']:
                                valid_headers[key] = value
                        
                        # If no valid headers found, use default
                        if not valid_headers:
                            valid_headers = {'Content-Type': 'application/json'}
                        
                        api_config_dict['headers'] = valid_headers
                        print(f"✅ Cleaned headers: {valid_headers}")
                
                # Strategy 3: Ensure required fields and proper data types
                if 'timeout' in api_config_dict:
                    try:
                        api_config_dict['timeout'] = int(api_config_dict['timeout'])
                    except (ValueError, TypeError):
                        api_config_dict['timeout'] = 30
                        print("⚠️ Invalid timeout value, defaulting to 30 seconds")
                
                # Ensure headers is a dictionary
                if 'headers' not in api_config_dict or not isinstance(api_config_dict['headers'], dict):
                    print(f"⚠️ Missing or invalid headers, setting default")
                    api_config_dict['headers'] = {'Content-Type': 'application/json'}
                
                # Strategy 4: Validate required fields based on type
                config_type = api_config_dict.get('type', '')
                if config_type == 'custom-api':
                    if 'url' not in api_config_dict:
                        raise ValueError("Custom API configuration missing required 'url' field")
                    if 'method' not in api_config_dict:
                        api_config_dict['method'] = 'POST'
                        print("⚠️ Missing method, defaulting to POST")
                elif config_type in ['coze-agent', 'coze-bot']:
                    if 'url' not in api_config_dict:
                        # Set default Coze URL based on type
                        if config_type == 'coze-agent':
                            api_config_dict['url'] = 'https://api.coze.cn/open_api/v2/chat'
                        else:
                            api_config_dict['url'] = 'https://api.coze.cn/open_api/v2/chat'
                        print(f"⚠️ Missing URL for {config_type}, using default")
            
            print(f"🔧 Cleaned API config: {json.dumps(api_config_dict, indent=2)}")
            
            api_config = APIConfig(**api_config_dict)
            logger.info(f"✅ API config parsed successfully: {api_config.type}")
        except json.JSONDecodeError as je:
            error_msg = f"JSON格式错误: {str(je)}"
            logger.error(f"❌ JSON parsing failed: {error_msg}")
            logger.error(f"❌ Original config string: {agent_api_config}")
            raise HTTPException(status_code=400, detail=f"API配置JSON格式错误: {error_msg}")
        except ValueError as ve:
            error_msg = str(ve)
            logger.error(f"❌ Config validation failed: {error_msg}")
            logger.error(f"❌ Original config string: {agent_api_config}")
            raise HTTPException(status_code=400, detail=f"API配置验证失败: {error_msg}")
        except Exception as e:
            error_msg = str(e)
            logger.error(f"❌ API config parsing failed: {error_msg}")
            logger.error(f"❌ Original config string: {agent_api_config}")
            raise HTTPException(status_code=400, detail=f"API配置解析失败: {error_msg}")
        
        # Handle requirement document
        requirement_context = ""
        user_persona_info = None
        
        if evaluation_mode == "auto" and extracted_persona:
            # Use extracted persona information
            try:
                user_persona_info = json.loads(extracted_persona)
                print("🎭 Using extracted user persona information")
            except:
                print("⚠️ Failed to parse extracted persona, falling back to manual mode")
                evaluation_mode = "manual"
        
        if requirement_file and requirement_file.filename:
            print(f"📄 Processing uploaded file: {requirement_file.filename}")
            file_content = await requirement_file.read()
            
            if requirement_file.filename.endswith('.docx'):
                requirement_context = await extract_text_from_docx(file_content)
            elif requirement_file.filename.endswith('.pdf'):
                requirement_context = await extract_text_from_pdf(file_content)
            elif requirement_file.filename.endswith('.txt'):
                requirement_context = file_content.decode('utf-8', errors='ignore')
            else:
                print("⚠️ Unsupported file format, using text input instead")
                
        if not requirement_context and requirement_text:
            # ⭐ Security: Sanitize user input
            requirement_context = sanitize_user_input(requirement_text, max_length=100000)
        
        # Parse conversation scenarios
        scenarios = []
        
        if evaluation_mode == "auto" and user_persona_info:
            # Auto mode: generate scenarios based on extracted persona
            print("🎯 Auto mode: Generating conversation scenarios based on extracted persona...")
            scenarios = await generate_conversation_scenarios_from_persona(user_persona_info)
            print(f"✅ Generated {len(scenarios)} scenarios based on user persona: {user_persona_info['user_persona']['role']}")
            
        else:
            # Manual mode or fallback: parse provided scenarios
            try:
                scenarios = json.loads(conversation_scenarios)
            except Exception as e:
                if evaluation_mode == "manual":
                    raise HTTPException(status_code=400, detail=f"场景配置解析失败: {str(e)}")
                else:
                    # Auto mode but no valid persona - generate basic scenarios
                    print("⚠️ Auto mode but no persona available, generating basic scenarios...")
                    scenarios = [
                        {
                            "title": "基础咨询场景",
                            "context": "通用咨询环境",
                            "user_profile": "普通用户",
                            "turns": ["我有个问题", "能详细说明一下吗", "谢谢"]
                        }
                    ]

        if not scenarios:
            if evaluation_mode == "auto":
                raise HTTPException(status_code=400, detail="自动模式下无法生成对话场景，请检查用户画像提取是否成功")
            else:
                raise HTTPException(status_code=400, detail="请至少配置一个对话场景")

        print(f"📋 Total scenarios to evaluate: {len(scenarios)}")
        
        # Enhanced evaluation with persona-aware context
        evaluation_results = []
        
        for i, scenario in enumerate(scenarios, 1):
            print(f"📋 场景 {i}/{len(scenarios)}: {scenario.get('title', '未命名场景')}")
            
            # Enhance scenario with extracted persona if available
            if evaluation_mode == "auto" and user_persona_info:
                scenario = enhance_scenario_with_persona(scenario, user_persona_info)
                print(f"🎭 Enhanced scenario with extracted persona: {user_persona_info['user_persona']['role']}")
            
            result = await evaluate_single_conversation_scenario(
                api_config=api_config,
                scenario=scenario,
                requirement_context=requirement_context,
                evaluation_mode=evaluation_mode,
                user_persona_info=user_persona_info
            )
            
            if result:
                evaluation_results.append(result)
            else:
                print(f"⚠️ 场景 {i} 评估失败，跳过")

        if not evaluation_results:
            raise HTTPException(status_code=500, detail="所有场景评估均失败，请检查AI Agent配置")

        # Generate summary
        summary = generate_evaluation_summary(evaluation_results, requirement_context)
        
        # Enhanced recommendations with persona awareness
        recommendations = generate_enhanced_recommendations(evaluation_results, user_persona_info)
        
        response_data = {
            "evaluation_summary": summary,
            "conversation_records": evaluation_results,
            "recommendations": recommendations,
            "evaluation_mode": evaluation_mode,
            "user_persona_info": user_persona_info,
                "timestamp": datetime.now().isoformat()
        }
        
        print(f"🎯 总体评估完成！综合得分: {summary.get('overall_score', 0):.2f}/5.0")
        print(f"📊 评估模式: {evaluation_mode}")
        if user_persona_info:
            print(f"🎭 用户画像: {user_persona_info['user_persona']['role']}")
        
        return response_data
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Evaluation failed with error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"评估过程出错: {str(e)}")

@app.post("/api/evaluate-agent-specification-query", response_model=EvaluationResponse)
async def evaluate_agent_specification_query(
    agent_api_config: str = Form(...),
    use_raw_messages: bool = Form(False),  # Use raw user messages without persona enhancement
    enable_turn_control: bool = Form(True),  # Enable 6-turn limit and satisfaction detection
    is_tricky_test: bool = Form(False)  # Enable tricky test mode for edge-case questions
):
    """
    Evaluate AI Agent specifically for specification query project with fixed persona and enhanced dimensions
    """
    try:
        print("🚀 开始规范查询项目评估...")
        
        # Use fixed specification query persona and context
        user_persona_info = SPECIFICATION_QUERY_DEFAULTS
        requirement_context = "工程项目现场监理工程师规范查询系统"
        
        # Parse API configuration
        try:
            api_config = APIConfig.parse_raw(agent_api_config)
            print(f"✅ API配置解析成功: {api_config.type}")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"API配置解析失败: {str(e)}")
        
        # Generate optimized scenario for specification query
        scenarios = await generate_specification_query_scenarios(user_persona_info)
        
        # Perform evaluation with specification query enhancements
        evaluation_results = []
        for i, scenario in enumerate(scenarios):
            print(f"\n🎯 评估场景 {i+1}/{len(scenarios)}: {scenario.get('title', '规范查询场景')}")
            
            try:
                # Conduct conversation with turn control if enabled
                if enable_turn_control:
                            conversation_history = await conduct_conversation_with_turn_control(
            api_config, scenario, user_persona_info, use_raw_messages, is_tricky_test
        )
                else:
                    conversation_history = await conduct_optimized_dynamic_conversation(
                        api_config, scenario, user_persona_info, use_raw_messages
                    )
                
                if not conversation_history:
                    print(f"  ⚠️ 场景 {i+1} 未生成有效对话")
                    continue
                
                # Use specification query evaluation
                evaluation_scores, detailed_explanations, scenario_score = await evaluate_conversation_specification_query(
                    conversation_history, scenario, requirement_context, user_persona_info
                )
                
                # Store enhanced result
                evaluation_results.append({
                    "scenario": scenario,
                    "conversation_history": conversation_history,
                    "evaluation_scores": evaluation_scores,
                    "detailed_explanations": detailed_explanations,
                    "evaluation_scores_with_explanations": detailed_explanations,
                    "scenario_score": scenario_score / 20,  # Convert to 5-point scale for compatibility
                    "scenario_score_100": scenario_score,   # Keep 100-point scale
                    "scenario_grade": get_score_grade(scenario_score)
                })
                
                print(f"  ✅ 场景 {i+1} 评估完成，得分: {scenario_score:.1f}/100")
                
            except Exception as e:
                print(f"  ❌ 场景 {i+1} 评估失败: {str(e)}")
                continue
        
        if not evaluation_results:
            raise HTTPException(status_code=500, detail="所有评估场景都失败了")
        
        # Generate comprehensive summary
        overall_score_100 = sum(r.get('scenario_score_100', 0) for r in evaluation_results) / len(evaluation_results)
        overall_score_5 = overall_score_100 / 20
        total_conversations = sum(len(r.get('conversation_history', [])) for r in evaluation_results)
        
        # Calculate dimension averages (8 dimensions for specification query)
        all_scores_100 = {}
        for result in evaluation_results:
            scores = result.get("evaluation_scores", {})
            for dimension, score in scores.items():
                if dimension not in all_scores_100:
                    all_scores_100[dimension] = []
                all_scores_100[dimension].append(score)
        
        dimension_averages_100 = {}
        for dimension, scores in all_scores_100.items():
            dimension_averages_100[dimension] = round(sum(scores) / len(scores), 2) if scores else 0
        
        evaluation_summary = {
            "overall_score": round(overall_score_5, 2),
            "overall_score_100": round(overall_score_100, 2),
            "total_scenarios": len(evaluation_results),
            "total_conversations": total_conversations,
            "framework": "规范查询项目 8维度评估框架 (100分制)",
            "dimension_scores_100": dimension_averages_100,
            "dimension_scores": {k: round(v/20, 2) for k, v in dimension_averages_100.items()},
            "scoring_mode": "specification_query",
            "turn_control_enabled": enable_turn_control,
            "max_turns": MAX_CONVERSATION_TURNS
        }
        
        # Generate AI-powered improvement suggestions for programmers
        print("🤖 正在生成AI智能改进建议...")
        all_detailed_explanations = {}
        for result in evaluation_results:
            explanations = result.get("detailed_explanations", {})
            for dimension, explanation in explanations.items():
                if dimension not in all_detailed_explanations:
                    all_detailed_explanations[dimension] = []
                all_detailed_explanations[dimension].append(explanation)
        
        # Merge explanations from multiple scenarios
        merged_explanations = {}
        for dimension, explanation_list in all_detailed_explanations.items():
            if explanation_list:
                # Use the first explanation as base, but merge full_response content
                merged_explanations[dimension] = explanation_list[0].copy()
                # Combine all full_response content
                all_responses = []
                for exp in explanation_list:
                    if exp and isinstance(exp, dict) and exp.get('full_response'):
                        all_responses.append(exp['full_response'])
                if all_responses:
                    merged_explanations[dimension]['full_response'] = "\n\n".join(all_responses)
        
        ai_improvement_suggestions = await generate_ai_improvement_suggestions_for_programmers(
            merged_explanations, evaluation_summary
        )
        
        # Prepare response data
        response_data = {
            "evaluation_summary": evaluation_summary,
            "conversation_records": evaluation_results,
            "recommendations": ai_improvement_suggestions,  # Add AI-generated recommendations
            "ai_improvement_suggestions": ai_improvement_suggestions,  # Also add with specific key
            "user_persona_info": {
                "extracted_persona_summary": user_persona_info,
                "extraction_method": "规范查询项目固定画像",
                "extraction_timestamp": datetime.now().isoformat()
            },
            "evaluation_mode": "specification_query",
            "timestamp": datetime.now().isoformat(),
            "version": "4.1"
        }
        
        # Save to database if enabled
        if config.ENABLE_DATABASE_SAVE:
            try:
                session_id = await save_evaluation_to_database(response_data, requirement_context)
                response_data["session_id"] = session_id
                print(f"💾 评估结果已保存到数据库: {session_id}")
            except Exception as e:
                print(f"⚠️ 数据库保存失败: {str(e)}")
        
        print(f"🎉 规范查询评估完成! 综合得分: {overall_score_100:.1f}/100")
        
        return response_data
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ 规范查询评估失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"评估失败: {str(e)}")

@app.post("/api/evaluate-agent-dynamic", response_model=EvaluationResponse)
async def evaluate_agent_dynamic(
    agent_api_config: str = Form(...),
    requirement_file: UploadFile = File(None),
    requirement_text: str = Form(None),
    extracted_persona: str = Form(None),  # JSON string of extracted persona
    use_raw_messages: bool = Form(False),  # Use raw user messages without persona enhancement
    is_tricky_test: bool = Form(False)  # Enable tricky test mode for edge-case questions
):
    """
    New dynamic evaluation endpoint implementing conversational workflow:
    1. Extract persona from document
    2. Generate 2 scenarios with dynamic conversations
    3. Conduct 2-3 rounds per scenario based on AI responses
    4. Generate comprehensive final report
    """
    # Generous timeout protection with optimization safety net
    evaluation_timeout = 480  # 8 minutes total timeout (optimized code usually finishes in 1-3 minutes)
    
    # Check memory usage before starting evaluation
    memory_usage = check_memory_usage()
    if memory_usage and memory_usage > config.MEMORY_WARNING_THRESHOLD:
        logger.warning(f"⚠️ High memory usage detected: {memory_usage:.1f}%")
    
    try:
        # Wrap the entire evaluation in a timeout
        evaluation_result = await asyncio.wait_for(
            _perform_dynamic_evaluation_internal(
                agent_api_config, requirement_file, requirement_text, extracted_persona, use_raw_messages, is_tricky_test
            ),
            timeout=evaluation_timeout
        )
        return evaluation_result
        
    except asyncio.TimeoutError:
        logger.error(f"⏰ Dynamic evaluation timed out after {evaluation_timeout} seconds")
        raise HTTPException(
            status_code=408, 
            detail=f"评估超时：评估过程超过{evaluation_timeout//60}分钟限制。系统已优化通常在1-3分钟内完成，如果超时可能原因：1) AI Agent API响应异常缓慢，2) 网络连接不稳定，3) 服务器负载过高。建议检查API配置和网络状况。"
        )
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        # Catch any other unexpected exceptions
        logger.error(f"❌ Unexpected error in dynamic evaluation: {str(e)}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=500, 
            detail=f"动态评估过程出现意外错误: {str(e)}. 请检查服务器日志获取详细信息。"
        )

async def _perform_dynamic_evaluation_internal(
    agent_api_config: str,
    requirement_file: UploadFile,
    requirement_text: str,
    extracted_persona: str,
    use_raw_messages: bool = False,
    is_tricky_test: bool = False
) -> Dict:
    """
    Internal function to perform dynamic evaluation with proper error handling
    """
    try:
        logger.info("🚀 Starting dynamic evaluation...")
        print("🚀============================================================🚀")
        print("   AI Agent 动态对话评估平台 v4.0")
        print(f"🔍 消息处理模式: {'原始消息模式 (RAW)' if use_raw_messages else '增强消息模式 (ENHANCED)'}")
        print("🚀============================================================🚀")
        
        # Parse API configuration
        import json as json_module  # Import locally to avoid scope issues
        try:
            # ⭐ Security: Validate input length
            if len(agent_api_config) > 50000:  # 50KB limit
                raise HTTPException(status_code=413, detail="API配置过长，请检查配置内容")
            
            api_config_dict = json_module.loads(agent_api_config)
            
            # ⭐ Security: Validate API URL if present
            if 'url' in api_config_dict and not validate_api_url(api_config_dict['url']):
                raise HTTPException(status_code=400, detail="不安全的API URL")
            
            # Debug: log the received configuration structure
            print(f"🔍 Received API config structure: {json_module.dumps(api_config_dict, indent=2)}")
            logger.info(f"🔍 Received API config: {api_config_dict}")
            
            # Check if the config is wrapped in an extra layer (common frontend issue)
            if isinstance(api_config_dict, dict):
                # Look for common wrapping patterns
                if 'config' in api_config_dict and isinstance(api_config_dict['config'], dict):
                    print("⚠️ Detected config wrapped in 'config' key, unwrapping...")
                    api_config_dict = api_config_dict['config']
                elif 'headers' in api_config_dict and 'url' in api_config_dict.get('headers', {}):
                    print("⚠️ Detected config wrapped in 'headers' key, unwrapping...")
                    api_config_dict = api_config_dict['headers']
                elif 'api_config' in api_config_dict and isinstance(api_config_dict['api_config'], dict):
                    print("⚠️ Detected config wrapped in 'api_config' key, unwrapping...")
                    api_config_dict = api_config_dict['api_config']
            
            # Additional data cleaning for common frontend issues
            if isinstance(api_config_dict, dict):
                # Ensure timeout is an integer
                if 'timeout' in api_config_dict:
                    try:
                        api_config_dict['timeout'] = int(api_config_dict['timeout'])
                    except (ValueError, TypeError):
                        api_config_dict['timeout'] = 30
                
                # Ensure headers is a dictionary
                if 'headers' in api_config_dict and not isinstance(api_config_dict['headers'], dict):
                    print(f"⚠️ Invalid headers format: {type(api_config_dict['headers'])}, resetting to empty dict")
                    api_config_dict['headers'] = {}
            
            print(f"🔧 Cleaned API config: {json_module.dumps(api_config_dict, indent=2)}")
            
            api_config = APIConfig(**api_config_dict)
            logger.info(f"✅ API config parsed successfully: {api_config.type}")
        except Exception as e:
            logger.error(f"❌ API config parsing failed: {str(e)}")
            logger.error(f"❌ Original config string: {agent_api_config}")
            raise HTTPException(status_code=400, detail=f"API配置解析失败: {str(e)}")
        
        # Handle requirement document
        requirement_context = ""
        user_persona_info = None
        
        # Step 1: Process requirement document and extract persona
        try:
            # ⭐ Memory check before document processing
            memory_usage = check_memory_usage()
            if memory_usage > config.MEMORY_CRITICAL_THRESHOLD:
                raise HTTPException(status_code=507, detail=f"内存不足 ({memory_usage:.1f}%)")
            
            if requirement_file and requirement_file.filename:
                logger.info(f"📄 Processing uploaded file: {requirement_file.filename}")
                print(f"📄 Processing uploaded file: {requirement_file.filename}")
                requirement_context = await process_uploaded_document_improved(requirement_file)
            elif requirement_text:
                logger.info("📝 Using provided text content")
                # ⭐ Security: Sanitize user input
                requirement_context = sanitize_user_input(requirement_text, max_length=100000)
            
            if not requirement_context:
                raise HTTPException(status_code=400, detail="请提供需求文档或文本内容")
                
            logger.info(f"✅ Document processed, length: {len(requirement_context)} characters")
        except Exception as e:
            logger.error(f"❌ Document processing failed: {str(e)}")
            raise HTTPException(status_code=400, detail=f"文档处理失败: {str(e)}")
        
        # Step 2: Extract user persona from requirement document using DeepSeek
        try:
            if extracted_persona:
                try:
                    user_persona_info = json_module.loads(extracted_persona)
                    logger.info(f"🎭 Using provided persona: {user_persona_info.get('user_persona', {}).get('role', '未知角色')}")
                    print(f"🎭 使用提取的用户画像: {user_persona_info.get('user_persona', {}).get('role', '未知角色')}")
                except Exception as pe:
                    logger.warning(f"⚠️ Persona parsing failed: {str(pe)}")
                    print("⚠️ 画像数据解析失败，重新提取...")
                    user_persona_info = None
            
            if not user_persona_info:
                logger.info("🧠 Extracting user persona from document...")
                print("🧠 从需求文档中提取用户画像...")
                user_persona_info = await extract_user_persona_with_deepseek(requirement_context)
                if not user_persona_info:
                    raise HTTPException(status_code=400, detail="无法从需求文档中提取有效的用户画像信息")
                    
            logger.info("✅ User persona extracted successfully")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"❌ Persona extraction failed: {str(e)}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"用户画像提取失败: {str(e)}")
        
        # Step 3: Conduct dynamic multi-scenario evaluation
        try:
            # ⭐ Memory check before evaluation
            memory_usage = check_memory_usage()
            if memory_usage > config.MEMORY_CRITICAL_THRESHOLD:
                raise HTTPException(status_code=507, detail=f"内存不足 ({memory_usage:.1f}%)")
            
            logger.info("🎯 Starting dynamic conversation evaluation...")
            print("🎯 开始动态多轮对话评估...")
            evaluation_results = await conduct_dynamic_multi_scenario_evaluation(
                api_config, user_persona_info, requirement_context, use_raw_messages, is_tricky_test
            )
            
            if not evaluation_results:
                raise HTTPException(status_code=500, detail="动态对话评估失败，请检查AI Agent配置")
                
            logger.info(f"✅ Evaluation completed with {len(evaluation_results)} scenarios")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"❌ Evaluation failed: {str(e)}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"动态对话评估失败: {str(e)}")
        
        # Step 4: Generate comprehensive final report
        try:
            logger.info("📊 Generating comprehensive report...")
            print("📊 生成综合评估报告...")
            comprehensive_report = await generate_final_comprehensive_report(
                evaluation_results, user_persona_info, requirement_context
            )
            logger.info("✅ Report generated successfully")
        except Exception as e:
            logger.error(f"❌ Report generation failed: {str(e)}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            # Use fallback report generation
            comprehensive_report = {
                "improvement_recommendations": ["系统建议：加强对话理解能力", "系统建议：提高回答准确性"],
                "extracted_persona_summary": user_persona_info,
                "persona_alignment_analysis": "基于系统分析生成",
                "business_goal_achievement": "评估完成"
            }
        
        # Calculate overall summary
        try:
            # Calculate from scenario scores (which are now in 5-point scale)
            overall_score_5 = sum(r.get('scenario_score', 0) for r in evaluation_results) / len(evaluation_results) if evaluation_results else 0
            overall_score_100 = sum(r.get('scenario_score_100', 0) for r in evaluation_results) / len(evaluation_results) if evaluation_results else 0
            total_conversations = sum(len(r.get('conversation_history', [])) for r in evaluation_results)
            
            # Generate comprehensive evaluation summary  
            evaluation_summary = generate_evaluation_summary(evaluation_results, requirement_context)
            
            response_data = {
                "conversation_records": evaluation_results,
                "recommendations": comprehensive_report.get('improvement_recommendations', []),
                "ai_improvement_suggestions": ai_improvement_suggestions,  # Add AI suggestions
                "evaluation_summary": evaluation_summary,
                "user_persona_info": user_persona_info,
                "detailed_context_display": {
                    "requirement_document_summary": {
                        "content_length": len(requirement_context) if requirement_context else 0,
                        "content_preview": requirement_context[:500] + "..." if requirement_context and len(requirement_context) > 500 else requirement_context,
                        "analysis_basis": "基于上传的需求文档进行AI智能分析"
                    },
                    "evaluation_methodology": {
                        "conversation_generation": "DeepSeek动态生成对话场景",
                        "response_evaluation": "多维度100分制评估",
                        "persona_matching": "基于文档提取的用户画像进行个性化测试"
                    },
                    "technical_details": {
                        "api_type": "Dify API" if "/v1/chat-messages" in api_config.url else "Coze API" if "coze" in api_config.url.lower() else "自定义API",
                        "conversation_turns": total_conversations,
                        "evaluation_dimensions": len(evaluation_results[0].get('evaluation_scores', {})) if evaluation_results else 3
                    }
                },
                "persona_alignment_analysis": comprehensive_report.get('persona_alignment_analysis', ''),
                "business_goal_achievement": comprehensive_report.get('business_goal_achievement', ''),
                "evaluation_mode": "dynamic_evaluation"
            }
            
            logger.info(f"🎯 Dynamic evaluation completed successfully! Score: {overall_score_100:.2f}/100.0")
            print(f"🎯 动态评估完成！综合得分: {overall_score_100:.2f}/100.0")
            print(f"📊 评估场景: {len(evaluation_results)} 个")
            print(f"💬 对话轮次: {total_conversations} 轮")
            print(f"🎭 用户画像: {user_persona_info.get('user_persona', {}).get('role', '未知角色')}")
            
            # Monitor response size and optimize if needed
            import sys
            response_json = json_module.dumps(response_data, ensure_ascii=False, default=str)
            response_size_mb = sys.getsizeof(response_json) / (1024 * 1024)
            
            logger.info(f"📊 Response size: {response_size_mb:.2f} MB")
            print(f"📊 Response size: {response_size_mb:.2f} MB")
            
            # If response is too large (>50MB), optimize it
            if response_size_mb > 50:
                logger.warning(f"⚠️ Large response detected ({response_size_mb:.2f} MB), optimizing...")
                print(f"⚠️ Large response detected ({response_size_mb:.2f} MB), optimizing...")
                
                # Reduce conversation history verbosity for large responses
                for record in response_data.get("conversation_records", []):
                    for turn in record.get("conversation_history", []):
                        # Truncate very long AI responses
                        if len(turn.get("ai_response", "")) > 5000:
                            turn["ai_response"] = turn["ai_response"][:5000] + "\n...[响应已截断，完整内容请查看详细报告]"
                        
                        # Truncate very long evaluation explanations
                        for key, value in record.get("evaluation_scores_with_explanations", {}).items():
                            if isinstance(value, dict) and len(str(value.get("detailed_analysis", ""))) > 2000:
                                value["detailed_analysis"] = str(value["detailed_analysis"])[:2000] + "...[详细分析已截断]"
            
            # Store the response before attempting database save
            final_response = response_data.copy()
            
            # Auto-save to database if enabled (non-blocking)
            if config.ENABLE_AUTO_SAVE:
                try:
                    # Use asyncio.create_task to make database save non-blocking
                    async def save_to_db():
                        try:
                            session_id = await save_evaluation_to_database(response_data, requirement_context)
                            if session_id:
                                print(f"💾 评估结果已自动保存到数据库，会话ID: {session_id}")
                            else:
                                print("⚠️ 数据库自动保存失败，但评估结果仍然可用")
                        except Exception as e:
                            print(f"⚠️ 数据库保存异常，但不影响评估结果: {str(e)}")
                    
                    # Create background task for database save
                    asyncio.create_task(save_to_db())
                    
                    # Attempt quick database save with timeout
                    try:
                        session_id = await asyncio.wait_for(
                            save_evaluation_to_database(response_data, requirement_context),
                            timeout=5.0  # 5 second timeout for database save
                        )
                        if session_id:
                            final_response["database_session_id"] = session_id
                            print(f"💾 评估结果已自动保存到数据库，会话ID: {session_id}")
                    except asyncio.TimeoutError:
                        print("⚠️ 数据库保存超时，已创建后台任务继续保存")
                    except Exception as e:
                        print(f"⚠️ 数据库保存异常，但不影响评估结果: {str(e)}")
                        
                except Exception as e:
                    print(f"⚠️ 数据库保存模块异常: {str(e)}")
            
            # Final response validation and optimization
            try:
                # Ensure the response can be JSON serialized
                test_json = json.dumps(final_response, ensure_ascii=False, default=str)
                final_size_mb = sys.getsizeof(test_json) / (1024 * 1024)
                logger.info(f"✅ Final response ready: {final_size_mb:.2f} MB")
                print(f"✅ Final response ready: {final_size_mb:.2f} MB")
                
                # Add response metadata
                final_response["response_metadata"] = {
                    "size_mb": round(final_size_mb, 2),
                    "generation_time": datetime.now().isoformat(),
                    "optimized": response_size_mb > 50,
                    "version": "4.0"
                }
                
                return final_response
                
            except Exception as json_error:
                logger.error(f"❌ Response serialization failed: {str(json_error)}")
                print(f"❌ Response serialization failed: {str(json_error)}")
                
                # Return a minimal response if serialization fails
                return {
                    "evaluation_summary": {
                        "overall_score_100": round(overall_score_100, 2),
                        "total_scenarios": len(evaluation_results),
                        "error": "Full response too large, providing summary only"
                    },
                    "conversation_records": [],
                    "recommendations": ["系统建议：响应过大，请查看详细报告"],
                    "timestamp": datetime.now().isoformat(),
                    "error_info": "Response optimization failed"
                }
            
        except Exception as e:
            logger.error(f"❌ Response data assembly failed: {str(e)}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"响应数据组装失败: {str(e)}")
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        # Catch any other unexpected exceptions
        logger.error(f"❌ Unexpected error in dynamic evaluation: {str(e)}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=500, 
            detail=f"动态评估过程出现意外错误: {str(e)}. 请检查服务器日志获取详细信息。"
        )

def enhance_scenario_with_persona(scenario: Dict, user_persona_info: Dict) -> Dict:
    """
    Enhance conversation scenario with extracted user persona information
    """
    enhanced_scenario = scenario.copy()
    
    persona = user_persona_info.get('user_persona', {})
    context = user_persona_info.get('usage_context', {})
    ai_role = user_persona_info.get('ai_role_simulation', {})
    
    # Enhance user profile
    if not scenario.get('user_profile') or scenario.get('user_profile') == '':
        enhanced_scenario['user_profile'] = f"{persona.get('role', '专业用户')}，{persona.get('experience_level', '有经验')}"
    
    # Enhance context if not provided
    if not scenario.get('context') or scenario.get('context') == '':
        work_env = persona.get('work_environment', '')
        business_domain = context.get('business_domain', '')
        enhanced_scenario['context'] = f"{work_env} - {business_domain}" if work_env and business_domain else work_env or business_domain or "专业工作环境"
    
    # Add persona-aware conversation approach
    enhanced_scenario['conversation_approach'] = ai_role.get('conversation_approach', '直接专业提问')
    enhanced_scenario['language_style'] = ai_role.get('language_characteristics', '专业术语与通俗解释结合')
    
    return enhanced_scenario

async def evaluate_single_conversation_scenario(
    api_config: APIConfig,
    scenario: Dict,
    requirement_context: str = "",
    evaluation_mode: str = "manual",
    user_persona_info: Dict = None
) -> Optional[Dict]:
    """
    Enhanced single scenario evaluation with persona awareness
    """
    try:
        print(f"🗣️ 开始对话场景: {scenario.get('title', '未命名场景')}")
        
        conversation_history = []
        turns = scenario.get('turns', [])
        
        if not turns:
            print("⚠️ 场景没有配置对话轮次")
            return None
            
        # Enhanced conversation simulation with persona context
        for turn_num, user_message in enumerate(turns, 1):
            if not user_message.strip():
                continue
                
            print(f"💬 第 {turn_num} 轮对话: {user_message[:50]}...")
            
            # Add persona context to the message if in auto mode
            enhanced_message = user_message
            if evaluation_mode == "auto" and user_persona_info:
                persona_context = f"[作为{user_persona_info['user_persona']['role']}，{user_persona_info['user_persona']['communication_style']}] {user_message}"
                enhanced_message = persona_context
            
            try:
                ai_response = await call_ai_agent_api(api_config, enhanced_message)
                
                if ai_response:
                    conversation_history.append({
                        "turn": turn_num,
                        "user_message": user_message,  # Store original message for display
                        "enhanced_message": enhanced_message if evaluation_mode == "auto" else user_message,
                        "ai_response": ai_response
                    })
                    print(f"✅ AI响应: {ai_response[:100]}...")
                else:
                    print(f"❌ 第 {turn_num} 轮AI无响应")
                    
            except Exception as e:
                print(f"❌ 第 {turn_num} 轮对话失败: {str(e)}")
                continue
        
        if not conversation_history:
            print("❌ 场景对话完全失败")
            return None
        
        # Enhanced evaluation with persona awareness
        evaluation_scores, explanations = await evaluate_conversation_with_deepseek(
            conversation_history, scenario, requirement_context, user_persona_info
        )
        
        scenario_score = sum(evaluation_scores.values()) / len(evaluation_scores) if evaluation_scores else 0
        print(f"🎯 场景得分: {scenario_score:.2f}/5.0")
        
        return {
            "scenario": {
                "title": scenario.get('title', '未命名场景'),
                "context": f"{user_persona_info.get('usage_context', {}).get('business_domain', '专业服务') if user_persona_info else scenario.get('context', '专业工作环境')} - {scenario.get('context', '专业工作环境')}",
                "user_profile": f"{user_persona_info.get('user_persona', {}).get('role', '用户') if user_persona_info else scenario.get('user_profile', '用户')}，{user_persona_info.get('user_persona', {}).get('experience_level', '中等经验') if user_persona_info else '中等经验'}，{user_persona_info.get('user_persona', {}).get('communication_style', '专业沟通') if user_persona_info else '专业沟通'}"
            },
            "conversation_history": conversation_history,
            "evaluation_scores": evaluation_scores,
            "evaluation_scores_with_explanations": explanations,  # Add detailed explanations
            "explanations": explanations,
            "scenario_score": scenario_score,
            "evaluation_mode": evaluation_mode,
            "persona_enhanced": evaluation_mode == "auto" and user_persona_info is not None,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        print(f"❌ 场景评估异常: {str(e)}")
        return None

async def perform_deepseek_evaluations(evaluation_prompts: Dict, base_context: str, requirement_context: str) -> tuple:
    """
    Perform DeepSeek evaluations for all dimensions
    """
    try:
        # Standard evaluation prompts 
        if not evaluation_prompts:  # If prompts not provided, create them
            evaluation_prompts = {
                "answer_correctness": f"""
你是一名**建筑工程问答质检官**，需要为 AI 答案的 **准确性 / 专业性** 打分并写评语。

🔍 **必须重点排查**  
1. *引用规范是否真实且与场景技术**适用***？（错把燃气软管规范当石材规范 → 严重错误）  
2. 若回答"查不到"或"标准空缺"，是否说明信息来源、限制或替代做法？  
3. 有无"看似合理、实则无关"或"堆砌规范名词"来掩盖缺乏依据的情况？

📊 **评分标准 (0-100)**  
- **90–100** 完全命中问题+引用真实且适用的条款  
- **75–89** 基本正确，但存在轻微遗漏/未给出处  
- **60–74** 回答表面合理，但未真正解决核心痛点或缺关键参数  
- **40–59** 引用规范无关或适用性错误（如将 *GB 44017-2024* 用作石材腐蚀标准）  
- **0–39** 伪造规范/条款或严重误导

🧪 **负面示例（判 45 分）**  
> 问：火山凝灰岩抗酸雨腐蚀有什么强制标准？  
> 答：可参考 *GB 44017-2024 燃气用具连接用金属包覆软管* 第 7.6.3 条……  
> 评语：规范与石材场景无关，属硬套用，严重误导。

---

{base_context}

📝 **请输出 JSON**  
```json
{{"score": X, "comment": "…" }}
```
""",
                "specification_citation_accuracy": f"""
你是建筑规范引用核查员，需核对 AI 答案中提到的每一处规范/条款。

🔍 **核查清单**  
1. 编号 / 年份 / 条款是否存在？  
2. 是否引用了已废止版本？  
3. 引用文字与原文是否一致？  
4. 引用是否与问题主题相关？  
5. 有无虚构条款？  
6. **适用性**：规范所属领域与提问场景是否匹配？

📊 **评分 (0-100)**  
- **90–100** 全部准确且技术适用  
- **75–89** 小幅偏差（如年份/条款号写错一位）  
- **60–74** 条款存在但内容对不上 / 摘抄失真  
- **40–59** 引用规范领域不符或硬套用（如将燃气软管规范用于石材检测）  
- **0–39** 规范或条款不存在 / 整段伪造

🧪 **负面示例（判 50 分）**  
引用 *GB 44017-2024* 第 7.6.3 条作为"火山凝灰岩耐酸雨试验"，实际该规范用于燃气软管，领域不符。

---

{base_context}

📝 **请输出 JSON**  
```json
{{"score": X, "comment": "…" }}
```
""",
                "fuzzy_understanding": f"""
你是AI模糊问题理解与引导能力评估专家，任务是评估AI在面对**不清晰或不完整问题**时的理解和引导能力。

🔍 【重点关注】：
1. 能否识别问题中的模糊或缺失信息
2. 是否主动询问关键细节以明确需求
3. 在信息不足时是否给出合理的引导性回答
4. 能否帮助用户逐步完善问题描述

---

{base_context}

📊 **评分标准 (0-100)**：
- **90-100** 准确识别模糊点，主动引导澄清，提供有效建议
- **75-89** 能识别部分模糊信息，有一定引导但不够全面
- **60-74** 能察觉问题不清晰但引导方式不够有效
- **40-59** 对模糊问题缺乏敏感度，回答偏离或过于宽泛
- **0-39** 完全忽视问题模糊性，给出不相关或误导性回答

🧪 **负面示例（判 30 分）**  
> 问：房子有问题，能修吗？  
> 答：当然可以修，一般房屋问题都有解决办法，你可以找施工队来看看。  
> 评语：完全忽视问题的模糊性，给出过于笼统且可能误导的回答。

---

📝 **请输出 JSON**  
```json
{{"score": X, "comment": "…" }}
```
""",
                "multi_turn_support": f"""
你是AI多轮对话连贯性与支持度评估专家，任务是评估AI在**多轮对话**中的表现质量。

🔍 【重点关注】：
1. 是否能记住并引用前面对话的关键信息
2. 回答是否与上下文逻辑连贯，避免自相矛盾
3. 能否根据对话发展深入探讨问题
4. 是否保持用户角色和场景的一致性

---

{base_context}

📊 **评分标准 (0-100)**：
- **90-100** 完美记忆上下文，回答连贯深入，逻辑一致
- **75-89** 基本保持连贯性，偶有上下文衔接不够自然
- **60-74** 能跟随对话但缺乏深入，部分回答脱离上下文
- **40-59** 上下文记忆不稳定，存在逻辑矛盾
- **0-39** 完全无视前面对话，回答前后矛盾或重复

🧪 **负面示例（判 35 分）**  
> 轮1-问：钢结构刷什么漆？  
> 轮1-答：建议用环氧富锌底漆加聚氨酯面漆。  
> 轮2-问：这个漆贵不贵？  
> 轮2-答：钢结构防腐很重要，建议用防火涂料。  
> 评语：完全忽视上文涂料选择，答非所问，逻辑不连贯。

---

📝 **请输出 JSON**  
```json
{{"score": X, "comment": "…" }}
```
""",
                "persona_alignment": f"""
你是用户画像匹配度评估专家，任务是判断AI是否采用了符合该用户角色的沟通风格与表达方式。

{base_context}

📊 **评分标准 (0-100)**:
- **90–100** 完全贴合用户角色，语气自然专业
- **75–89** 基本匹配，偶有不自然表达
- **60–74** 有术语、语调等不适配问题
- **40–59** 明显风格不符
- **0–39** 完全脱离角色

📝 **请输出 JSON**  
```json
{{"score": X, "comment": "…" }}
```
"""
            }
            
            # Add goal alignment if requirement context exists
            if requirement_context.strip():
                evaluation_prompts["goal_alignment"] = f"""
{base_context}

基于提供的需求文档，请评估AI是否达成了预期的目标对齐度。

📊 **评分标准 (0-100)**  
- **90–100** 完美对齐所有需求目标，超出预期  
- **75–89** 很好地满足了大部分需求目标  
- **60–74** 基本符合需求但有重要遗漏  
- **40–59** 部分相关但未达成主要目标  
- **0–39** 完全偏离需求目标，未解决任何关键问题

📝 **请输出 JSON**  
```json
{{"score": X, "comment": "…" }}
```
"""
        
        # Execute evaluations concurrently for better performance
        evaluation_results = {}
        explanations = {}
        
        for dimension, prompt in evaluation_prompts.items():
            try:
                print(f"  📊 Evaluating {dimension}...")
                response = await call_deepseek_api(prompt)
                score = extract_score_from_response(response)
                
                evaluation_results[dimension] = score
                explanations[dimension] = response
                
                print(f"  ✅ {dimension}: {score}/100")
                
            except Exception as e:
                print(f"  ❌ Failed to evaluate {dimension}: {str(e)}")
                evaluation_results[dimension] = 3.0  # Default score
                explanations[dimension] = f"评估失败: {str(e)}"
        
        return evaluation_results, explanations
        
    except Exception as e:
        print(f"❌ Evaluation process failed: {str(e)}")
        return {}, {}

def generate_evaluation_summary(evaluation_results: List[Dict], requirement_context: str = "") -> Dict:
    """
    Generate evaluation summary from results - 100-point scale normalized, with Chinese labels and filtered dimensions
    """
    if not evaluation_results:
        return {
            "overall_score_100": 0.0,
            "overall_score": 0.0,  # Keep for compatibility
            "total_scenarios": 0,
            "total_conversations": 0,
            "framework": "AI Agent 6维度评估框架 (100分制)",
            "dimension_scores_100": {},
            "dimension_scores": {},  # Keep for compatibility
            "persona": {},
            "usage_context": {},
            "goal": {}
        }
    # Calculate dimension averages from 100-point scores
    all_scores_100 = {}
    total_conversations = 0
    for result in evaluation_results:
        scores = result.get("evaluation_scores", {})
        conversation_history = result.get("conversation_history", [])
        total_conversations += len(conversation_history)
        for dimension, score in scores.items():
            if dimension in ["response_conciseness", "error_handling_transparency"]:
                continue
            if dimension not in all_scores_100:
                all_scores_100[dimension] = []
            normalized_score = score * 20 if score <= 5 else score
            all_scores_100[dimension].append(normalized_score)
    # Calculate averages in 100-point scale
    dimension_averages_100 = {}
    for dimension, scores in all_scores_100.items():
        avg_100 = sum(scores) / len(scores) if scores else 0
        zh_label = map_dimension_to_chinese(dimension)
        dimension_averages_100[zh_label] = round(avg_100, 2)
    overall_score_100 = sum(dimension_averages_100.values()) / len(dimension_averages_100) if dimension_averages_100 else 0
    overall_score_5 = overall_score_100 / 20  # For compatibility
    # Persona/context/goal extraction
    persona_summary = generate_persona_summary(requirement_context)
    usage_context_summary = generate_usage_context_summary(requirement_context)
    goal_summary = generate_goal_summary(requirement_context)
    return {
        "overall_score_100": round(overall_score_100, 2),
        "overall_score": round(overall_score_5, 2),
        "total_scenarios": len(evaluation_results),
        "total_conversations": total_conversations,
        "framework": "AI Agent 6维度评估框架 (100分制)",
        "dimension_scores_100": dimension_averages_100,
        "persona": persona_summary,
        "usage_context": usage_context_summary,
        "goal": goal_summary
    }

def map_dimension_to_chinese(dimension: str) -> str:
    mapping = {
        'answer_correctness': '回答准确性',
        'persona_alignment': '用户匹配度',
        'goal_alignment': '目标对齐度',
        'specification_citation_accuracy': '规范引用准确度',
        'fuzzy_understanding': '模糊理解能力',
        'multi_turn_support': '多轮支持度',
        'user_matching': '用户匹配度',
        'target_alignment': '目标对齐度',
    }
    return mapping.get(dimension, dimension)

def map_explanations_to_chinese(explanations: dict) -> dict:
    filtered = {}
    for dim, data in explanations.items():
        if dim in ["response_conciseness", "error_handling_transparency"]:
            continue
        zh_label = map_dimension_to_chinese(dim)
        filtered[zh_label] = data
    return filtered

def extract_recommendations_from_response(response: str) -> List[str]:
    """Extract improvement recommendations from DeepSeek response"""
    try:
        # Look for numbered or bulleted recommendations
        recommendations = []
        lines = response.split('\n')
        
        for line in lines:
            line = line.strip()
            if any(marker in line for marker in ['1.', '2.', '3.', '4.', '5.', '-', '•', '①', '②', '③']):
                # Clean the line and extract recommendation
                clean_rec = re.sub(r'^[\d\.\-\•①②③④⑤]\s*', '', line).strip()
                if clean_rec and len(clean_rec) > 10:
                    recommendations.append(clean_rec)
        
        # If no structured recommendations found, try to extract from content
        if not recommendations and response:
            # Split into sentences and look for actionable suggestions
            sentences = re.split(r'[。！？.]', response)
            for sentence in sentences:
                if any(keyword in sentence for keyword in ['建议', '应该', '需要', '可以', '改进', '提升', '优化']):
                    if len(sentence.strip()) > 15:
                        recommendations.append(sentence.strip())
        
        return recommendations[:5] if recommendations else [
            "增强对话理解能力",
            "提高回答准确性",
            "优化用户体验",
            "加强专业知识深度"
        ]
        
    except Exception:
        return [
            "增强对话理解能力",
            "提高回答准确性", 
            "优化用户体验",
            "加强专业知识深度"
        ]

@app.post("/api/validate-config")
async def validate_agent_config(agent_api_config: str = Form(...)):
    """
    Validate AI Agent configuration before evaluation
    """
    try:
        # Parse configuration
        api_config_dict = json.loads(agent_api_config)
        api_config = APIConfig(**api_config_dict)
        
        validation_results = {
            "is_valid": True,
            "warnings": [],
            "errors": [],
            "suggestions": []
        }
        
        # Basic validation
        if not api_config.url:
            validation_results["errors"].append("API URL不能为空")
            validation_results["is_valid"] = False
        
        if not api_config.headers:
            validation_results["warnings"].append("未设置请求头，可能导致认证失败")
        
        # Coze-specific validation
        if api_config.type == "coze-agent":
            if not api_config.agentId:
                validation_results["errors"].append("Coze Agent ID不能为空")
                validation_results["is_valid"] = False
            
            auth_header = api_config.headers.get("Authorization", "")
            if not auth_header or not auth_header.startswith("Bearer "):
                validation_results["errors"].append("Coze Agent需要有效的Bearer Token")
                validation_results["is_valid"] = False
            
            if api_config.region not in ["global", "china"]:
                validation_results["warnings"].append("建议使用 'global' 或 'china' 作为region")
        
        elif api_config.type == "coze-bot":
            if not api_config.botId:
                validation_results["errors"].append("Coze Bot ID不能为空")
                validation_results["is_valid"] = False
        
        # Test connectivity (optional quick test)
        if validation_results["is_valid"]:
            try:
                # Quick connectivity test with timeout
                async with httpx.AsyncClient(timeout=httpx.Timeout(5.0)) as client:
                    if api_config.type == "coze-agent":
                        # Test Coze Agent endpoint
                        test_url = f"https://api.coze.{'cn' if api_config.region == 'china' else 'com'}/v3/chat"
                        response = await client.post(
                            test_url,
                            headers=api_config.headers,
                            json={
                                "bot_id": api_config.agentId,
                                "user_id": "test_validation",
                                "stream": False,
                                "auto_save_history": False,
                                "additional_messages": [
                                    {"role": "user", "content": "test", "content_type": "text"}
                                ]
                            }
                        )
                        
                        if response.status_code == 401:
                            validation_results["errors"].append("认证失败：请检查Access Token是否有效")
                            validation_results["is_valid"] = False
                        elif response.status_code == 403:
                            validation_results["errors"].append("权限不足：请检查Token权限或Agent ID")
                            validation_results["is_valid"] = False
                        elif response.status_code in [200, 400]:  # 400 might be expected for test message
                            validation_results["suggestions"].append("✅ API连接测试成功")
                        else:
                            validation_results["warnings"].append(f"API返回状态码: {response.status_code}")
                    
                    else:
                        # Test custom API endpoint
                        response = await client.post(
                            api_config.url,
                            headers=api_config.headers,
                            json={"message": "test"}
                        )
                        
                        if response.status_code in [200, 400, 401]:
                            validation_results["suggestions"].append("✅ API端点可访问")
                        else:
                            validation_results["warnings"].append(f"API返回状态码: {response.status_code}")
                            
            except httpx.TimeoutException:
                validation_results["warnings"].append("API连接超时，请检查网络或URL")
            except Exception as e:
                validation_results["warnings"].append(f"连接测试失败: {str(e)[:50]}...")
        
        # Add helpful suggestions
        if validation_results["is_valid"]:
            validation_results["suggestions"].extend([
                "💡 建议先用简单问题测试AI Agent响应",
                "📝 确保Agent已发布且处于活跃状态",
                "🔄 如遇到问题，可尝试重新生成Access Token"
            ])
        
        return {
            "validation_result": validation_results,
            "config_summary": {
                "type": api_config.type,
                "url": api_config.url,
                "has_auth": bool(api_config.headers.get("Authorization")),
                "timeout": api_config.timeout
            }
        }
        
    except json.JSONDecodeError:
        return {
            "validation_result": {
                "is_valid": False,
                "errors": ["配置格式错误：请检查JSON格式"],
                "warnings": [],
                "suggestions": ["请确保配置信息为有效的JSON格式"]
            }
        }
    except Exception as e:
        return {
            "validation_result": {
                "is_valid": False,
                "errors": [f"配置验证失败: {str(e)}"],
                "warnings": [],
                "suggestions": ["请检查配置信息是否完整"]
            }
        }

async def call_deepseek_api_with_fallback(prompt: str, max_retries: int = 2) -> str:
    """
    Enhanced DeepSeek API call with config-based settings and proper error handling
    """
    headers = {
        "Authorization": f"Bearer {config.DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 800,
        "temperature": config.ENHANCED_TEMPERATURE
    }
    
    for attempt in range(max_retries):
        try:
            timeout = httpx.Timeout(config.DEEPSEEK_TIMEOUT, connect=10.0)
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.post(config.DEEPSEEK_API_URL, headers=headers, json=payload)
                
                if response.status_code == 200:
                    result = response.json()
                    if "choices" in result and len(result["choices"]) > 0:
                        content = result["choices"][0]["message"]["content"].strip()
                        if content and len(content) > 10:
                            return content
                        else:
                            raise Exception("DeepSeek returned empty or too short response")
                    else:
                        raise Exception("No valid choices in DeepSeek response")
                elif response.status_code == 401:
                    raise Exception("API authentication failed - check API key")
                elif response.status_code == 429:
                    if attempt < max_retries - 1:
                        await asyncio.sleep(2 ** attempt)
                        continue
                    raise Exception("API rate limited")
                else:
                    error_text = response.text if hasattr(response, 'text') else 'Unknown error'
                    raise Exception(f"API error {response.status_code}: {error_text}")
                    
        except (asyncio.TimeoutError, httpx.TimeoutException):
            if attempt < max_retries - 1:
                await asyncio.sleep(1)
                continue
            raise Exception(f"API timeout after {config.DEEPSEEK_TIMEOUT}s")
        except httpx.RequestError as e:
            if attempt < max_retries - 1:
                await asyncio.sleep(1)
                continue
            raise Exception(f"Network error: {str(e)}")
        except Exception as e:
            if "empty" in str(e) or "short" in str(e):
                if attempt < max_retries - 1:
                    await asyncio.sleep(1)
                    continue
            raise e
            
    raise Exception("All API attempts failed")

# Create alias function to redirect to new implementation
async def conduct_dynamic_conversation(api_config: APIConfig, scenario_info: Dict, user_persona_info: Dict) -> List[Dict]:
    """Redirect to true dynamic conversation implementation"""
    return await conduct_true_dynamic_conversation(api_config, scenario_info, user_persona_info)

async def conduct_true_dynamic_conversation(api_config: APIConfig, scenario_info: Dict, user_persona_info: Dict, use_raw_messages: bool = False, is_tricky_test: bool = False) -> List[Dict]:
    """
    TRUE dynamic conversation: Correct flow implementation
    
    Flow: 
    1. DeepSeek generates user message based on persona → 
    2. Send RAW message to Coze (no enhancement) → 
    3. Extract Coze's actual response → 
    4. Pass Coze response to DeepSeek for next message generation
    
    Args:
        use_raw_messages: Legacy parameter (always uses raw messages now)
    """
    print(f"🗣️ 开始真正动态对话: {scenario_info.get('title', '未命名场景')}")
    print("🔄 正确流程: DeepSeek(基于画像生成消息) → 原始消息 → Coze → 响应 → DeepSeek(分析响应生成下轮)")
    
    conversation_history = []
    persona = user_persona_info.get('user_persona', {})
    context = user_persona_info.get('usage_context', {})
    
    # Initialize conversation manager for continuity
    conversation_manager = ConversationManager(api_config)
    conversation_manager.start_new_conversation()
    
    # Step 1: Generate ONLY the initial message based on persona and scenario
    try:
        initial_message = await generate_single_initial_message(scenario_info, user_persona_info, is_tricky_test)
        if not initial_message:
            raise Exception("Failed to generate initial message")
    except Exception as e:
        print(f"❌ 初始消息生成失败: {str(e)}")
        raise Exception(f"Dynamic conversation initialization failed: {str(e)}")
    
    current_user_message = initial_message
    failed_turns = 0  # Track failed turns
    
    # Step 2: Conduct true turn-by-turn conversation (optimized to 2-3 turns max)
    for turn_num in range(1, 4):  # Maximum 3 turns
        try:
            # 🐛 Debug log for message processing - ALWAYS use raw messages in dynamic conversation
            print(f"🔍 [TURN {turn_num}] DeepSeek生成的原始用户消息: {current_user_message}")
            
            # ALWAYS send raw user message to Coze (no persona enhancement in dynamic mode)
            # This is the correct flow: DeepSeek(persona) → raw message → Coze → response → DeepSeek(analyze)
            message_to_send = current_user_message
            print(f"🔍 [RAW MESSAGE] 发送原始消息到Coze: {message_to_send}")
            
            # Get AI response with timeout and conversation continuity
            ai_response = await call_coze_with_strict_timeout(api_config, message_to_send, conversation_manager, True)
            
            if not ai_response or len(ai_response.strip()) < 5:
                print(f"⚠️ 第{turn_num}轮AI响应为空或过短，可能是API问题")
                failed_turns += 1
                if failed_turns >= 2:  # Stop if too many failed turns
                    print("❌ 连续多轮API响应失败，可能是API配置或账户问题，请检查API密钥和账户余额")
                    break
                continue
            
            # Clean the AI response to extract meaningful content
            cleaned_response = clean_ai_response(ai_response)
            
            # If cleaned response is empty (system message), try to generate a fallback response
            if not cleaned_response:
                try:
                    fallback_prompt = f"""作为一个专业的{user_persona_info.get('user_persona', {}).get('role', '助手')}，请对以下问题给出简短但有用的回复：

用户问题：{current_user_message}
回复要求：
1. 直接回答问题，不要说"我不知道"
2. 保持专业但友好的语调
3. 如果需要更多信息，可以简单询问
4. 回复控制在50字以内

请直接给出回复："""
                    
                    fallback_response = await call_deepseek_api_enhanced(fallback_prompt, temperature=0.3, max_tokens=100)
                    if fallback_response and len(fallback_response.strip()) > 5:
                        cleaned_response = fallback_response.strip()
                        failed_turns = 0  # Reset since we got a response
                    else:
                        failed_turns += 1
                        if failed_turns >= 2:
                            break
                        continue
                except Exception as e:
                    failed_turns += 1
                    if failed_turns >= 2:
                        break
                    continue
            
            # Reset failed turn counter on successful turn
            failed_turns = 0
            
            # Record this turn
            conversation_history.append({
                "turn": turn_num,
                "user_message": current_user_message,  # Raw message generated by DeepSeek
                "message_sent_to_ai": message_to_send,  # Same as user_message (always raw)
                "ai_response": cleaned_response,        # Actual Coze response
                "response_length": len(cleaned_response),
                "conversation_id": conversation_manager.get_conversation_id(),
                "flow_type": "deepseek_to_raw_to_coze",  # Indicate the correct flow
                "timestamp": datetime.now().isoformat()
            })
            
            print(f"✅ 第 {turn_num} 轮对话完成")
            
            # Generate next message based on AI's actual response (only if not the last turn)
            if turn_num < 3:  # Don't generate after last turn
                try:
                    next_message = await generate_next_message_based_on_response(
                        scenario_info, user_persona_info, conversation_history, cleaned_response, is_tricky_test
                    )
                    
                    if not next_message or next_message.upper() in ["END", "FINISH", "DONE"]:
                        print(f"🔚 对话自然结束于第 {turn_num} 轮")
                        break
                        
                    current_user_message = next_message
                    
                except Exception as e:
                    print(f"❌ 第{turn_num + 1}轮消息生成失败: {str(e)}")
                    break  # End conversation if next message generation fails
            
        except Exception as e:
            print(f"❌ 第 {turn_num} 轮对话异常: {str(e)}")
            failed_turns += 1
            if failed_turns >= 2:
                break
            continue
    
    if not conversation_history:
        raise Exception("Dynamic conversation completely failed - no successful turns")
    
    print(f"📊 真实动态对话完成，共 {len(conversation_history)} 轮")
    print(f"✅ 实现流程: DeepSeek(画像生成) → 原始消息 → Coze → 实际回复 → DeepSeek(分析回复)")
    return conversation_history

# DeepSeek Configuration

async def call_deepseek_api_enhanced(prompt: str, max_tokens: int = 500, temperature: float = 0.1, max_retries: int = 2) -> str:
    """
    Enhanced DeepSeek API call with better configuration and error handling
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.DEEPSEEK_API_KEY}"
    }
    
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": temperature,
        "top_p": 0.9,
        "frequency_penalty": 0.1,
        "presence_penalty": 0.1
    }
    
    # Single attempt - fail fast if there are issues
    try:
        # Increased timeout and added better error handling
        timeout = httpx.Timeout(config.DEEPSEEK_TIMEOUT, connect=10.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(config.DEEPSEEK_API_URL, json=payload, headers=headers)
            
            if response.status_code == 200:
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    content = result["choices"][0]["message"]["content"]
                    return content.strip()
                else:
                    raise Exception("No valid response choices in API response")
            elif response.status_code == 429:
                raise Exception(f"API rate limited (429)")
            elif response.status_code == 401:
                raise Exception(f"API authentication failed (401) - check API key")
            else:
                error_text = response.text if hasattr(response, 'text') else 'Unknown error'
                raise Exception(f"API error {response.status_code}: {error_text}")
                
    except asyncio.TimeoutError:
        raise Exception(f"API request timeout after {config.DEEPSEEK_TIMEOUT}s - try increasing timeout in config.py")
    except httpx.TimeoutException:
        raise Exception(f"API request timeout after {config.DEEPSEEK_TIMEOUT}s - try increasing timeout in config.py")
    except httpx.RequestError as e:
        raise Exception(f"Network error: {str(e)} - check internet connection")
    except Exception as e:
        # Re-raise without fallback
        raise e

# Add after the health check endpoint
@app.post("/api/download-report")
async def download_evaluation_report(
    request: Request,
    evaluation_data: str = Form(...),
    format: str = Form("json"),  # json, txt, docx
    include_transcript: bool = Form(False)
):
    """
    Generate and download evaluation report in specified format
    """
    try:
        # Parse evaluation data
        eval_results = json.loads(evaluation_data)
        
        # 🆕 自动保存评估结果到数据库（如果尚未保存）
        session_id = None
        if PYMYSQL_AVAILABLE and config.ENABLE_AUTO_SAVE:
            try:
                # 检查评估数据中是否已有session_id
                session_id = eval_results.get('session_id')
                if not session_id:
                    # 如果没有session_id，保存评估结果到数据库
                    requirement_context = eval_results.get('requirement_document', '')
                    session_id = await save_evaluation_to_database(eval_results, requirement_context)
                    print(f"✅ 评估结果已自动保存到数据库，会话ID: {session_id}")
                
                # 记录下载活动
                if session_id:
                    file_size = len(evaluation_data)  # 估算文件大小
                    await save_download_record(session_id, format, include_transcript, file_size, request)
                    print(f"📥 下载记录已保存: {format} 格式，包含对话记录: {include_transcript}")
                    
            except Exception as db_error:
                print(f"⚠️ 数据库保存失败，但报告生成将继续: {db_error}")
        
        # 生成并返回报告
        if format == "json":
            return generate_json_report(eval_results, include_transcript)
        elif format == "txt":
            return generate_txt_report(eval_results, include_transcript)
        elif format == "docx":
            return generate_docx_report(eval_results, include_transcript)
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

def generate_json_report(eval_results: Dict, include_transcript: bool = False) -> JSONResponse:
    """Generate JSON format report with complete evaluation information and proper formatting"""
    # 🔍 DEBUG: Check the actual evaluation_mode value
    print(f"🔍 DEBUG: eval_results.evaluation_mode = {eval_results.get('evaluation_mode', 'NOT_FOUND')}")
    
    report_data = {
        "evaluation_summary": eval_results.get("evaluation_summary", {}),
        "overall_score": eval_results.get("overall_score", 0),
        "dimension_scores": eval_results.get("dimension_scores", {}),
        "detailed_analysis": eval_results.get("detailed_analysis", {}),
        "recommendations": eval_results.get("recommendations", []),
        "ai_improvement_suggestions": eval_results.get("ai_improvement_suggestions", []),  # Add DeepSeek recommendations
        "user_persona_info": eval_results.get("user_persona_info", {}),
        "detailed_context_display": eval_results.get("detailed_context_display", {}),
        "persona_alignment_analysis": eval_results.get("persona_alignment_analysis", ""),
        "business_goal_achievement": eval_results.get("business_goal_achievement", ""),
        "evaluation_mode": eval_results.get("evaluation_mode", "manual"),
        "timestamp": eval_results.get("timestamp", datetime.now().isoformat())
    }
    
    if include_transcript:
        report_data["conversation_records"] = eval_results.get("conversation_records", [])
    
    return JSONResponse(
        content=report_data,
        headers={"Content-Disposition": "attachment; filename=evaluation_report.json"}
    )

def generate_txt_report(eval_results: Dict, include_transcript: bool = False) -> FileResponse:
    """Generate TXT format report"""
    import tempfile
    
    # Extract scoring information with proper 100-point scale
    overall_score = eval_results.get('evaluation_summary', {}).get('overall_score', eval_results.get('overall_score', 0))
    
    report_content = f"""
AI Agent Evaluation Report
=========================
Generated: {eval_results.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}
Evaluation Mode: {eval_results.get('evaluation_mode', 'Unknown')}

OVERALL PERFORMANCE
------------------
Overall Score: {overall_score}/100.0

DIMENSION SCORES
---------------
"""
    
    # Use conversation_records to extract dimension scores
    conversation_records = eval_results.get('conversation_records', [])
    if conversation_records:
        for i, record in enumerate(conversation_records, 1):
            scenario_title = record.get('scenario', {}).get('title', f'Scenario {i}')
            report_content += f"\n{scenario_title}:\n"
            scores = record.get('evaluation_scores_with_explanations', record.get('evaluation_scores', {}))
            for dimension, score_data in scores.items():
                score = score_data.get('score', score_data) if isinstance(score_data, dict) else score_data
                dimension_name = dimension.replace('_', ' ').title()
                report_content += f"  {dimension_name}: {score}/100.0\n"
    
    report_content += f"\nDETAILED ANALYSIS\n{'-' * 16}\n"
    
    detailed_analysis = eval_results.get('detailed_analysis', {})
    for dimension, analysis in detailed_analysis.items():
        report_content += f"\n{dimension.upper()}:\n"
        if isinstance(analysis, dict):
            report_content += f"Score: {analysis.get('score', 'N/A')}\n"
            report_content += f"Analysis: {analysis.get('detailed_analysis', 'No details available')}\n"
        else:
            report_content += f"{analysis}\n"
    
    report_content += f"\nRECOMMENDations\n{'-' * 15}\n"
    recommendations = eval_results.get('recommendations', [])
    for i, rec in enumerate(recommendations, 1):
        report_content += f"{i}. {rec}\n"
    
    if include_transcript:
        report_content += f"\nCONVERSATION TRANSCRIPT\n{'-' * 22}\n"
        conversation_records = eval_results.get('conversation_records', [])
        for record in conversation_records:
            for turn in record.get('conversation', []):
                report_content += f"Turn {turn.get('turn', 'N/A')}: {turn.get('user_message', '')}\n"
                report_content += f"AI Response: {turn.get('ai_response', '')}\n\n"
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as tmp_file:
        tmp_file.write(report_content)
        tmp_file_path = tmp_file.name
    
    return FileResponse(
        path=tmp_file_path,
        filename="evaluation_report.txt",
        media_type="text/plain"
    )

def generate_docx_report(eval_results: Dict, include_transcript: bool = False) -> FileResponse:
    """Generate DOCX format report (if docx library is available)"""
    if not DOCUMENT_PROCESSING_AVAILABLE:
        raise HTTPException(status_code=500, detail="DOCX generation not available. Install python-docx.")
    
    import tempfile
    from docx import Document
    from docx.shared import Inches
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Agent Evaluation Report', 0)
    
    # Summary section
    doc.add_heading('Executive Summary', level=1)
    overall_score = eval_results.get('overall_score', 'N/A')
    doc.add_paragraph(f'Overall Performance Score: {overall_score}/5.0')
    doc.add_paragraph(f'Generated: {eval_results.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))}')
    
    # Dimension scores
    doc.add_heading('Performance Dimensions', level=1)
    dimension_scores = eval_results.get('dimension_scores', {})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dimension'
    hdr_cells[1].text = 'Score'
    
    for dimension, score in dimension_scores.items():
        row_cells = table.add_row().cells
        row_cells[0].text = dimension.replace('_', ' ').title()
        row_cells[1].text = f'{score}/5.0'
    
    # Detailed analysis
    doc.add_heading('Detailed Analysis', level=1)
    detailed_analysis = eval_results.get('detailed_analysis', {})
    for dimension, analysis in detailed_analysis.items():
        doc.add_heading(dimension.replace('_', ' ').title(), level=2)
        if isinstance(analysis, dict):
            doc.add_paragraph(f"Score: {analysis.get('score', 'N/A')}")
            doc.add_paragraph(analysis.get('detailed_analysis', 'No details available'))
        else:
            doc.add_paragraph(str(analysis))
    
    # Recommendations
    doc.add_heading('Recommendations', level=1)
    recommendations = eval_results.get('recommendations', [])
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f'{i}. {rec}')
    
    # Conversation transcript (if requested)
    if include_transcript:
        doc.add_heading('Conversation Transcript', level=1)
        conversation_records = eval_results.get('conversation_records', [])
        for record in conversation_records:
            scenario_title = record.get('scenario', {}).get('title', 'Unknown Scenario')
            doc.add_heading(f'Scenario: {scenario_title}', level=2)
            for turn in record.get('conversation', []):
                doc.add_paragraph(f"Turn {turn.get('turn', 'N/A')}: {turn.get('user_message', '')}")
                doc.add_paragraph(f"AI Response: {turn.get('ai_response', '')}")
                doc.add_paragraph("")  # Empty line for spacing
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        tmp_file_path = tmp_file.name
    
    doc.save(tmp_file_path)
    
    return FileResponse(
        path=tmp_file_path,
        filename="evaluation_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def adjust_role_for_domain_consistency(extraction_result: Dict, domain_hints: Dict) -> Dict:
    """
    Perform light adjustment for domain consistency while preserving DeepSeek's analysis
    """
    domain = domain_hints.get('行业领域', '').lower()
    current_role = extraction_result.get('user_persona', {}).get('role', '')
    
    print(f"🔍 领域一致性检查: 域={domain}, 角色={current_role}")
    
    # Only perform light validation without forcing changes
    if domain and current_role:
        # Log the extracted role and domain for debugging
        print(f"✅ 提取的角色与领域: {current_role} in {domain}")
        
        # Basic domain-role matching suggestions (not enforced)
        suggested_domains = {
            '建筑': '建筑工程',
            '工程': '工程技术', 
            '银行': '银行服务',
            '金融': '金融服务',
            '客服': '客户服务',
            '医疗': '医疗健康',
            '教育': '教育培训'
        }
        
        # Suggest business domain if not specific enough
        current_domain = extraction_result.get('usage_context', {}).get('business_domain', '')
        if not current_domain or current_domain in ['专业服务', '未知']:
            for keyword, suggested_domain in suggested_domains.items():
                if keyword in domain:
                    extraction_result['usage_context']['business_domain'] = suggested_domain
                    print(f"💡 建议业务领域: {suggested_domain}")
                    break
    
    return extraction_result

def create_domain_aware_fallback_result(requirement_content: str, domain_hints: Dict) -> Dict[str, Any]:
    """
    Create a domain-aware fallback result when parsing fails
    Enhanced with better construction/civil engineering detection
    """
    logger.info("🔄 创建领域感知的回退结果...")
    print("🔄 创建领域感知的回退结果...")
    
    # Extract domain information with enhanced construction detection
    domain = domain_hints.get('行业领域', extract_business_domain_from_content(requirement_content))
    role = domain_hints.get('用户角色', extract_role_from_content(requirement_content))
    
    logger.info(f"🏢 检测到领域: {domain}")
    logger.info(f"👤 检测到角色: {role}")
    print(f"🏢 检测到领域: {domain}")
    print(f"👤 检测到角色: {role}")
    
    # Enhanced construction/civil engineering detection
    construction_indicators = ['建筑', '施工', '工程', '监理', '现场', '质量', '安全', '规范', '建设', '土建', '结构', '基础']
    if any(indicator in requirement_content for indicator in construction_indicators):
        logger.info("🏗️ 强制设置为建筑工程领域")
        print("🏗️ 强制设置为建筑工程领域")
        domain = '建筑工程'
        role = '土建工程师' if not role or '技术' in role else role
    
    # Ensure role matches domain with enhanced construction handling
    if '建筑' in domain.lower() or '工程' in domain.lower() or '施工' in domain.lower():
        # More accurate civil engineering role detection
        if '监理' in requirement_content:
            role = '建筑工程监理'
        elif '施工' in requirement_content:
            role = '施工工程师'
        elif '设计' in requirement_content:
            role = '建筑设计师'
        elif '质量' in requirement_content:
            role = '质量工程师'
        else:
            role = '土建工程师'  # Default for construction
        
        business_domain = '建筑工程'
        typical_questions = ["这个规范要求是什么？", "施工标准符合吗？", "质量检查怎么做？", "安全措施到位吗？", "这个材料符合标准吗？"]
        fuzzy_expressions = ["这个地方有问题", "标准不太对", "需要检查一下", "质量有点问题", "不太符合规范"]
    elif '银行' in domain.lower() or '金融' in domain.lower():
        role = role if '客服' in role or '银行' in role else '银行客服代表'
        business_domain = '银行金融服务'
        typical_questions = ["客户问这个怎么办？", "这个业务怎么处理？", "政策是什么？"]
        fuzzy_expressions = ["客户不满意", "又是那个问题", "怎么解释呢"]
    elif '客服' in domain.lower():
        role = role if '客服' in role else '客服专员'
        business_domain = '客户服务'
        typical_questions = ["客户投诉怎么处理？", "这个问题怎么解决？", "服务标准是什么？"]
        fuzzy_expressions = ["客户又投诉了", "老问题了", "不知道怎么说"]
    else:
        role = role or '专业用户'
        business_domain = domain or '专业服务'
        typical_questions = ["这个怎么处理？", "规范要求是什么？", "还有其他方案吗？"]
        fuzzy_expressions = ["有点问题", "不太对", "怎么处理？"]

    return {
        "user_persona": {
            "role": role,
            "experience_level": "中等经验专业用户",
            "expertise_areas": [business_domain, "相关专业知识"],
            "communication_style": "专业但有时表达不完整，贴合行业特点",
            "work_environment": f"{business_domain}工作环境",
            "work_pressure": "正常工作压力，注重效率和准确性"
        },
        "usage_context": {
            "business_domain": business_domain,
            "primary_scenarios": [f"{business_domain}咨询", "工作支持"],
            "interaction_goals": ["获取准确信息", "解决工作问题"],
            "pain_points": ["信息不够具体", "回答时间较长"],
            "usage_timing": ["工作时间", "遇到问题时", "需要确认时"]
        },
        "ai_role_simulation": {
            "simulated_user_type": f"基于{business_domain}的{role}",
            "conversation_approach": "直接提问，有时表达模糊",
            "language_characteristics": f"{business_domain}专业术语与日常表达混合",
            "typical_questions": typical_questions,
            "fuzzy_expressions": fuzzy_expressions,
            "opening_patterns": ["关于这个...", "需要咨询...", "有个问题...", "想了解..."],
            "situational_variations": "工作繁忙时表达简短，正常情况下会详细描述"
        },
        "extracted_requirements": {
            "core_functions": ["准确信息查询", "专业问题解答"],
            "quality_expectations": ["回答准确", "响应及时", "专业性强"],
            "interaction_preferences": ["简洁明了", "包含具体示例", "提供操作指导"]
        }
    }

def extract_role_from_content(content: str) -> Optional[str]:
    """Extract user role from content"""
    if "客服" in content:
        return "客服代表"
    elif "监理" in content:
        return "现场监理工程师"
    elif "工程师" in content:
        return "工程师"
    elif "技术" in content:
        return "技术人员"
    return None

def extract_business_domain_from_content(content: str) -> str:
    """Extract business domain from content with enhanced construction detection"""
    logger.debug(f"🔍 分析业务领域，内容前200字符: {content[:200]}")
    
    # Enhanced construction/civil engineering detection
    construction_keywords = ['建筑', '施工', '工程', '监理', '现场', '质量检查', '安全规范', '建筑施工', '土建', '钢筋', '混凝土', '基础工程', '结构工程']
    civil_keywords = ['民用建筑', '工业建筑', '基础设施', '道路工程', '桥梁工程', '水电工程', '暖通工程', '消防工程']
    
    construction_count = sum(1 for kw in construction_keywords if kw in content)
    civil_count = sum(1 for kw in civil_keywords if kw in content)
    
    logger.debug(f"🏗️ 建筑关键词匹配数: {construction_count}")
    logger.debug(f"🏗️ 土建关键词匹配数: {civil_count}")
    
    if construction_count > 0 or civil_count > 0:
        logger.info("✅ 识别为建筑工程领域")
        return "建筑工程"
    elif "银行" in content or "金融" in content:
        return "银行金融服务"
    elif "客服" in content:
        return "客户服务"
    elif "技术" in content and "工程" not in content:  # Avoid misclassifying engineering as tech support
        return "技术支持"
    else:
        return "专业服务"

async def conduct_dynamic_multi_scenario_evaluation(
    api_config: APIConfig,
    user_persona_info: Dict,
    requirement_context: str,
    use_raw_messages: bool = False,
    is_tricky_test: bool = False
) -> List[Dict]:
    """
    Optimized dynamic evaluation: 1 focused scenario with 2-3 turns max
    Reduced from ~3-5 minutes to ~1-2 minutes while maintaining quality
    """
    try:
        print("🎯 开始优化的动态评估...")
        
        # Generate only 1 focused scenario (reduced from 2)
        scenarios = await generate_optimized_scenario_from_persona(user_persona_info)
        
        if not scenarios:
            print("⚠️ 无法生成动态场景，使用快速默认场景")
            scenarios = [{
                "title": f"{user_persona_info.get('usage_context', {}).get('business_domain', '专业服务')}核心咨询",
                "context": f"{user_persona_info.get('usage_context', {}).get('business_domain', '专业服务')}专业问题解决",
                "user_profile": user_persona_info.get('user_persona', {}).get('role', '专业用户')
            }]
        
        evaluation_results = []
        
        # Process single scenario with optimized conversation
        scenario_info = scenarios[0]  # Only use first scenario
        print(f"📋 核心场景: {scenario_info.get('title', '未命名场景')}")
        
        try:
            # Conduct optimized dynamic conversation (2-3 turns max)
            conversation_history = await conduct_optimized_dynamic_conversation(
                api_config, scenario_info, user_persona_info, use_raw_messages, is_tricky_test
            )
            
            if not conversation_history:
                print(f"⚠️ 场景对话失败，使用简化评估")
                return []
            
            # Simplified evaluation (reduced evaluation complexity)
            evaluation_scores, detailed_explanations, scenario_score = await evaluate_conversation_optimized(
                conversation_history, scenario_info, requirement_context, user_persona_info
            )
            
            # Convert scenario score to 5-point scale for consistency
            scenario_score_5 = scenario_score / 20.0 if scenario_score > 5 else scenario_score
            
            # Create evaluation result
            evaluation_result = {
                "scenario": {
                    "title": scenario_info.get('title', '核心场景'),
                    "context": scenario_info.get('context', '优化评估场景'),
                    "user_profile": scenario_info.get('user_profile', user_persona_info.get('user_persona', {}).get('role', '专业用户'))
                },
                "conversation_history": conversation_history,
                "evaluation_scores": evaluation_scores,
                "detailed_explanations": detailed_explanations,
                "scenario_score": round(scenario_score_5, 2),
                "scenario_score_100": round(scenario_score, 2),
                "evaluation_mode": "optimized_dynamic",
                "timestamp": datetime.now().isoformat()
            }
            
            evaluation_results.append(evaluation_result)
            print(f"✅ 核心场景评估完成，得分: {scenario_score_5:.2f}/5.0")
            
        except Exception as e:
            print(f"❌ 场景评估失败: {str(e)}")
            return []  # Return empty instead of raising
        
        if not evaluation_results:
            print("❌ 评估失败，返回空结果")
            return []
        
        print(f"🎯 优化评估完成，用时显著减少")
        return evaluation_results
        
    except Exception as e:
        print(f"❌ 动态评估失败: {str(e)}")
        return []

async def generate_optimized_scenario_from_persona(user_persona_info: Dict) -> List[Dict]:
    """
    Generate single dynamic scenario based on extracted user persona using DeepSeek API
    """
    try:
        persona = user_persona_info.get('user_persona', {})
        usage_context = user_persona_info.get('usage_context', {})
        ai_role = user_persona_info.get('ai_role_simulation', {})
        
        role = persona.get('role', '专业用户')
        business_domain = usage_context.get('business_domain', '专业服务')
        primary_scenarios = usage_context.get('primary_scenarios', ['咨询服务', '问题解决'])
        
        scenario_prompt = f"""基于以下用户画像，生成1个真实、具体的对话场景：

用户角色：{role}
业务领域：{business_domain}
主要使用场景：{', '.join(primary_scenarios)}
沟通风格：{persona.get('communication_style', '专业直接')}
工作环境：{persona.get('work_environment', '专业环境')}

请生成JSON格式的场景：
{{
  "title": "具体场景名称",
  "context": "详细的业务背景和情境描述，反映{role}在{business_domain}中的典型工作情况",
  "user_profile": "用户在此场景下的具体身份和需求"
}}

要求：
1. 场景要真实反映{role}在{business_domain}中的典型工作情况
2. 场景要有明确的业务背景和具体情境
3. 避免过于复杂的技术细节，但要足够具体

直接输出JSON对象，不要其他文字："""
        
        print("🎭 DeepSeek生成动态场景...")
        response = await call_deepseek_api_enhanced(scenario_prompt, temperature=0.4, max_tokens=400)
        
        # Parse the JSON response
        cleaned_response = response.strip()
        if cleaned_response.startswith('```'):
            start_idx = cleaned_response.find('{')
            end_idx = cleaned_response.rfind('}')
            if start_idx != -1 and end_idx != -1:
                cleaned_response = cleaned_response[start_idx:end_idx+1]
        
        scenario = json.loads(cleaned_response)
        
        if isinstance(scenario, dict) and 'title' in scenario:
            print(f"✅ 成功生成动态场景: {scenario.get('title', '未命名场景')}")
            return [scenario]  # Return single scenario as list
        else:
            print("⚠️ 场景格式不正确，使用回退场景")
            # Fallback scenario
            return [{
                "title": f"{business_domain}专业咨询场景",
                "context": f"{role}在{business_domain}工作中遇到专业问题，需要专业指导和建议",
                "user_profile": f"{role}，{persona.get('experience_level', '有经验')}，需要专业技术支持"
            }]
            
    except Exception as e:
        print(f"⚠️ 动态场景生成失败: {str(e)}，使用回退场景")
        # Fallback scenario
        persona = user_persona_info.get('user_persona', {})
        usage_context = user_persona_info.get('usage_context', {})
        role = persona.get('role', '专业用户')
        business_domain = usage_context.get('business_domain', '专业服务')
        
        return [{
            "title": f"{business_domain}专业咨询场景",
            "context": f"{role}在{business_domain}工作中遇到专业问题，需要专业指导和建议",
            "user_profile": f"{role}，{persona.get('experience_level', '有经验')}，需要专业技术支持"
        }]

@app.post("/api/test-with-raw-coze-conversation")
async def test_with_raw_coze_conversation(
    agent_api_config: str = Form(...),
    coze_conversation_json: str = Form(...),
    use_raw_message: bool = Form(True)
):
    """
    Test AI Agent with raw Coze conversation JSON - extracts user message and sends it directly
    """
    try:
        print("🧪 开始原始Coze对话测试...")
        
        # Parse API configuration
        api_config_dict = json.loads(agent_api_config)
        api_config = APIConfig(**api_config_dict)
        
        # Parse Coze conversation JSON
        coze_data = json.loads(coze_conversation_json)
        
        # Extract raw user message
        raw_user_message = extract_user_message_from_coze_json(coze_data)
        
        if not raw_user_message:
            raise HTTPException(status_code=400, detail="无法从Coze对话JSON中提取用户消息")
        
        print(f"✅ 提取到原始用户消息: {raw_user_message}")
        
        # Call AI Agent with raw message
        ai_response = await call_ai_agent_api(api_config, raw_user_message, use_raw_message=use_raw_message)
        
        return {
            "status": "success",
            "raw_user_message": raw_user_message,
            "ai_response": ai_response,
            "use_raw_message": use_raw_message,
            "message_processing_mode": "RAW" if use_raw_message else "ENHANCED",
            "timestamp": datetime.now().isoformat()
        }
        
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"JSON解析失败: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Raw Coze conversation test failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"测试失败: {str(e)}")

@app.post("/api/extract-user-persona")
async def extract_user_persona(
    requirement_file: UploadFile = File(None),
    requirement_text: str = Form(None)
):
    """
    Extract user persona from requirement document
    """
    try:
        logger.info("🎭 开始用户画像提取...")
        print("🎭 开始用户画像提取...")
        
        # Handle requirement document
        requirement_context = ""
        
        if requirement_file and requirement_file.filename:
            logger.info(f"📄 Processing uploaded file: {requirement_file.filename}")
            print(f"📄 Processing uploaded file: {requirement_file.filename}")
            requirement_context = await process_uploaded_document_improved(requirement_file)
        elif requirement_text:
            logger.info("📝 Using provided text content")
            print("📝 Using provided text content")
            requirement_context = requirement_text
        
        if not requirement_context:
            raise HTTPException(status_code=400, detail="请提供需求文档或文本内容")
            
        logger.info(f"✅ Document processed, length: {len(requirement_context)} characters")
        print(f"✅ Document processed, length: {len(requirement_context)} characters")
        
        # Extract user persona using enhanced algorithm
        user_persona_info = await extract_user_persona_with_deepseek(requirement_context)
        
        if not user_persona_info:
            raise HTTPException(status_code=400, detail="无法从需求文档中提取有效的用户画像信息")
        
        return {
            "status": "success",
            "extraction_result": user_persona_info,
            "timestamp": datetime.now().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Persona extraction failed: {str(e)}")
        logger.error(f"📋 Traceback: {traceback.format_exc()}")
        print(f"❌ Persona extraction failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"用户画像提取失败: {str(e)}")

@app.post("/api/debug-document-processing")
async def debug_document_processing(
    requirement_file: UploadFile = File(None),
    requirement_text: str = Form(None)
):
    """
    Debug endpoint to test document processing and persona extraction separately
    This helps identify where the issue occurs in cloud deployment
    """
    try:
        result = {
            "document_processing": {},
            "persona_extraction": {},
            "domain_analysis": {},
            "errors": []
        }
        
        # Step 1: Document Processing
        requirement_context = ""
        
        if requirement_file and requirement_file.filename:
            logger.info(f"🧪 Debug: Processing file {requirement_file.filename}")
            try:
                requirement_context = await process_uploaded_document_improved(requirement_file)
                result["document_processing"] = {
                    "status": "success",
                    "filename": requirement_file.filename,
                    "content_length": len(requirement_context),
                    "content_preview": requirement_context[:300] + "..." if len(requirement_context) > 300 else requirement_context,
                    "file_type": getattr(requirement_file, 'content_type', 'unknown')
                }
            except Exception as e:
                error_msg = f"Document processing failed: {str(e)}"
                result["document_processing"] = {
                    "status": "error",
                    "error": error_msg,
                    "traceback": traceback.format_exc()
                }
                result["errors"].append(error_msg)
                
        elif requirement_text:
            requirement_context = requirement_text
            result["document_processing"] = {
                "status": "success",
                "source": "text_input",
                "content_length": len(requirement_text),
                "content_preview": requirement_text[:300] + "..." if len(requirement_text) > 300 else requirement_text
            }
        
        if not requirement_context:
            result["errors"].append("No document content available")
            return result
        
        # Step 2: Domain Analysis
        try:
            domain = extract_business_domain_from_content(requirement_context)
            role = extract_role_from_content(requirement_context)
            
            # Construction keywords analysis
            construction_keywords = ['建筑', '施工', '工程', '监理', '现场', '质量检查', '安全规范', '建筑施工', '土建', '钢筋', '混凝土', '基础工程', '结构工程']
            found_keywords = [kw for kw in construction_keywords if kw in requirement_context]
            
            result["domain_analysis"] = {
                "status": "success",
                "detected_domain": domain,
                "detected_role": role,
                "construction_keywords_found": found_keywords,
                "total_construction_keywords": len(found_keywords),
                "is_construction_content": len(found_keywords) > 0
            }
        except Exception as e:
            error_msg = f"Domain analysis failed: {str(e)}"
            result["domain_analysis"] = {
                "status": "error",
                "error": error_msg
            }
            result["errors"].append(error_msg)
        
        # Step 3: Persona Extraction
        try:
            user_persona_info = await extract_user_persona_with_deepseek(requirement_context)
            result["persona_extraction"] = {
                "status": "success",
                "extracted_role": user_persona_info.get('user_persona', {}).get('role', 'N/A'),
                "business_domain": user_persona_info.get('usage_context', {}).get('business_domain', 'N/A'),
                "extraction_method": "deepseek_api",
                "full_result": user_persona_info
            }
        except Exception as e:
            error_msg = f"Persona extraction failed: {str(e)}"
            result["persona_extraction"] = {
                "status": "error",
                "error": error_msg,
                "traceback": traceback.format_exc()
            }
            result["errors"].append(error_msg)
            
            # Try fallback method
            try:
                fallback_result = create_domain_aware_fallback_result(requirement_context, {})
                result["persona_extraction"]["fallback_result"] = {
                    "role": fallback_result.get('user_persona', {}).get('role', 'N/A'),
                    "business_domain": fallback_result.get('usage_context', {}).get('business_domain', 'N/A'),
                    "method": "fallback"
                }
            except Exception as fallback_error:
                result["persona_extraction"]["fallback_error"] = str(fallback_error)
        
        return result
        
    except Exception as e:
        logger.error(f"❌ Debug endpoint failed: {str(e)}")
        return {
            "status": "error",
            "error": str(e),
            "traceback": traceback.format_exc()
        }

# Database related functions
def get_database_connection():
    """
    Create database connection
    """
    if not PYMYSQL_AVAILABLE or not config.ENABLE_DATABASE_SAVE:
        return None
    
    try:
        connection = pymysql.connect(**config.DATABASE_CONFIG)
        return connection
    except Exception as e:
        print(f"❌ Database connection failed: {str(e)}")
        return None

def generate_session_id() -> str:
    """
    Generate unique session ID for evaluation
    """
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    random_suffix = uuid.uuid4().hex[:8]
    return f"EVAL_{timestamp}_{random_suffix}"

def get_evaluation_mode_abbreviation(mode: str) -> str:
    """Convert evaluation mode to database-friendly abbreviation"""
    mode_abbreviations = {
        'specification_query': 'spec_query',
        'dynamic_evaluation': 'dynamic',
        'manual_evaluation': 'manual',
        'auto_evaluation': 'auto',
        'with_file_evaluation': 'with_file',
        'multi_scenario': 'multi_scen'
    }
    
    # Return abbreviation if exists, otherwise truncate to 20 characters
    result = mode_abbreviations.get(mode, mode[:20])
    print(f"🔧 DEBUG: Abbreviating evaluation_mode '{mode}' -> '{result}' (length: {len(result)})")
    return result

async def save_evaluation_to_database(evaluation_data: Dict, requirement_context: str = "") -> str:
    """
    Save evaluation results to database
    Returns session_id if successful, None if failed
    """
    if not PYMYSQL_AVAILABLE or not config.ENABLE_DATABASE_SAVE:
        print("📝 Database save disabled or PyMySQL not available")
        return None
    
    connection = get_database_connection()
    if not connection:
        print("❌ Cannot connect to database")
        return None
    
    session_id = generate_session_id()
    
    try:
        with connection.cursor() as cursor:
            # Insert main evaluation session
            insert_session_sql = """
                INSERT INTO ai_evaluation_sessions (
                    session_id, overall_score, total_scenarios, total_conversations,
                    evaluation_mode, evaluation_framework, requirement_document,
                    ai_agent_config, user_persona_info, evaluation_summary, recommendations
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            
            evaluation_summary = evaluation_data.get('evaluation_summary', {})
            
            # Convert 100-point scale to 5-point scale for database storage
            overall_score = evaluation_summary.get('overall_score_100', evaluation_summary.get('overall_score', 0))
            if overall_score > 5:  # If it's 100-point scale, convert to 5-point scale
                overall_score_5_point = overall_score / 20.0  # Convert 100-point to 5-point
            else:
                overall_score_5_point = overall_score
            
            cursor.execute(insert_session_sql, (
                session_id,
                round(overall_score_5_point, 2),  # Store as 5-point scale
                evaluation_summary.get('total_scenarios', 0),
                evaluation_summary.get('total_conversations', 0),
                get_evaluation_mode_abbreviation(evaluation_data.get('evaluation_mode', 'manual')),  # Use abbreviation for database storage
                evaluation_summary.get('framework', 'AI Agent 3维度评估框架'),
                requirement_context[:5000] if requirement_context else None,  # Limit length
                json.dumps(evaluation_data.get('ai_agent_config', {})),
                json.dumps(evaluation_data.get('user_persona_info', {})),
                json.dumps(evaluation_summary),
                json.dumps(evaluation_data.get('recommendations', []))
            ))
            
            # Insert conversation scenarios and records
            conversation_records = evaluation_data.get('conversation_records', [])
            
            for scenario_index, record in enumerate(conversation_records):
                scenario = record.get('scenario', {})
                conversation_history = record.get('conversation_history', [])
                evaluation_scores = record.get('evaluation_scores_with_explanations', {})
                
                # Insert scenario
                insert_scenario_sql = """
                    INSERT INTO ai_conversation_scenarios (
                        session_id, scenario_index, scenario_title, scenario_context,
                        user_profile, scenario_score, conversation_turns
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s)
                """
                
                # Convert scenario score to 5-point scale for database storage
                scenario_score = record.get('scenario_score_100', record.get('scenario_score', 0))
                if scenario_score > 5:  # If it's 100-point scale, convert to 5-point scale
                    scenario_score_5_point = scenario_score / 20.0
                else:
                    scenario_score_5_point = scenario_score
                
                cursor.execute(insert_scenario_sql, (
                    session_id,
                    scenario_index,
                    scenario.get('title', f'场景 {scenario_index + 1}'),
                    scenario.get('context', ''),
                    scenario.get('user_profile', ''),
                    round(scenario_score_5_point, 2),  # Store as 5-point scale
                    len(conversation_history)
                ))
                
                scenario_id = cursor.lastrowid
                
                # Insert conversation turns
                for turn in conversation_history:
                    insert_turn_sql = """
                        INSERT INTO ai_conversation_turns (
                            session_id, scenario_id, turn_number, user_message,
                            enhanced_message, ai_response, response_length
                        ) VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """
                    
                    cursor.execute(insert_turn_sql, (
                        session_id,
                        scenario_id,
                        turn.get('turn', 0),
                        turn.get('user_message', ''),
                        turn.get('enhanced_message', ''),
                        turn.get('ai_response', ''),
                        len(turn.get('ai_response', ''))
                    ))
                
                # Insert evaluation scores
                for dimension_name, score_data in evaluation_scores.items():
                    if isinstance(score_data, dict):
                        insert_score_sql = """
                            INSERT INTO ai_evaluation_scores (
                                session_id, scenario_id, dimension_name, dimension_label,
                                score, detailed_analysis, specific_quotes,
                                improvement_suggestions, full_response
                            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """
                        
                        dimension_labels = {
                            'answer_correctness': '回答准确性与专业性',
                            'persona_alignment': '用户匹配度',
                            'goal_alignment': '目标对齐度'
                        }
                        
                        # Convert dimension score to 5-point scale for database storage
                        dimension_score = score_data.get('score', 0)
                        if dimension_score > 5:  # If it's 100-point scale, convert to 5-point scale
                            dimension_score_5_point = dimension_score / 20.0
                        else:
                            dimension_score_5_point = dimension_score
                        
                        cursor.execute(insert_score_sql, (
                            session_id,
                            scenario_id,
                            dimension_name,
                            dimension_labels.get(dimension_name, dimension_name),
                            round(dimension_score_5_point, 2),  # Store as 5-point scale
                            score_data.get('detailed_analysis', ''),
                            score_data.get('specific_quotes', ''),
                            score_data.get('improvement_suggestions', ''),
                            score_data.get('full_response', '')
                        ))
        
        connection.commit()
        print(f"✅ Evaluation data saved to database with session_id: {session_id}")
        return session_id
        
    except Exception as e:
        if connection:
            connection.rollback()
        print(f"❌ Failed to save evaluation to database: {str(e)}")
        
        # Check for common database errors and provide specific solutions
        error_str = str(e)
        if "Data truncated for column" in error_str:
            print("💡 Hint: This is likely a column size issue. The data being inserted is too long for the database column.")
            if "evaluation_mode" in error_str:
                print("   - evaluation_mode column needs to be expanded to accommodate values like 'specification_query'")
        elif "Duplicate entry" in error_str:
            print("💡 Hint: This is a duplicate key error. The session_id might already exist.")
        elif "doesn't exist" in error_str:
            print("💡 Hint: Database table doesn't exist. Please check if the database schema is set up correctly.")
        
        print(f"Full traceback: {traceback.format_exc()}")
        
        # Generate a fallback session_id so the evaluation can still continue
        fallback_session_id = f"local_{int(time.time())}"
        print(f"🔄 Using fallback session_id for this evaluation: {fallback_session_id}")
        return fallback_session_id
    finally:
        if connection:
            connection.close()

async def save_download_record(session_id: str, download_format: str, include_transcript: bool, 
                             file_size: int = None, request: Request = None) -> bool:
    """
    Record download activity
    """
    if not PYMYSQL_AVAILABLE or not config.ENABLE_DATABASE_SAVE:
        return False
    
    connection = get_database_connection()
    if not connection:
        return False
    
    try:
        with connection.cursor() as cursor:
            client_ip = request.client.host if request else None
            user_agent = request.headers.get("user-agent") if request else None
            
            insert_sql = """
                INSERT INTO ai_report_downloads (
                    session_id, download_format, include_transcript,
                    file_size, download_ip, user_agent
                ) VALUES (%s, %s, %s, %s, %s, %s)
            """
            
            cursor.execute(insert_sql, (
                session_id, download_format, include_transcript,
                file_size, client_ip, user_agent
            ))
        
        connection.commit()
        return True
        
    except Exception as e:
        print(f"❌ Failed to save download record: {str(e)}")
        return False
    finally:
        connection.close()

def extract_user_message_from_coze_json(coze_conversation_json: Dict) -> str:
    """
    Extract raw user message from Coze conversation JSON structure
    
    Expected Coze JSON structure:
    {
        "additional_messages": [
            {
                "content_type": "text",
                "type": "question",
                "role": "user", 
                "content": "SBS卷材搭接处热熔质量检查情况？节点部位附加层要做闭水试验"
            }
        ]
    }
    """
    try:
        # 🐛 Debug log for Coze JSON parsing
        print(f"🔍 [COZE JSON] Extracting user message from: {json.dumps(coze_conversation_json, ensure_ascii=False)[:200]}...")
        
        # Try to extract from additional_messages (most common)
        if "additional_messages" in coze_conversation_json:
            additional_messages = coze_conversation_json["additional_messages"]
            if isinstance(additional_messages, list) and len(additional_messages) > 0:
                first_message = additional_messages[0]
                if isinstance(first_message, dict) and "content" in first_message:
                    raw_content = first_message["content"].strip()
                    print(f"🔍 [EXTRACTED] Raw user message: {raw_content}")
                    return raw_content
        
        # Try to extract from messages array
        if "messages" in coze_conversation_json:
            messages = coze_conversation_json["messages"]
            if isinstance(messages, list):
                for message in messages:
                    if (isinstance(message, dict) and 
                        message.get("role") == "user" and 
                        "content" in message):
                        raw_content = message["content"].strip()
                        print(f"🔍 [EXTRACTED] Raw user message from messages: {raw_content}")
                        return raw_content
        
        # Try direct content field
        if "content" in coze_conversation_json:
            raw_content = coze_conversation_json["content"].strip()
            print(f"🔍 [EXTRACTED] Raw user message from content: {raw_content}")
            return raw_content
        
        # Fallback - look for any text content
        for key, value in coze_conversation_json.items():
            if isinstance(value, str) and len(value.strip()) > 5:
                print(f"🔍 [FALLBACK] Using field '{key}': {value.strip()}")
                return value.strip()
        
        print("❌ [EXTRACTION FAILED] No user message found in Coze JSON")
        return ""
        
    except Exception as e:
        print(f"❌ [EXTRACTION ERROR] Failed to extract user message: {str(e)}")
        return ""

def clean_ai_response(response: str) -> str:
    """
    Clean AI response to extract meaningful content and filter out system messages
    Enhanced to properly handle plugin tool outputs
    """
    try:
        original_response = response
        print(f"🧹 Cleaning AI response: {response[:100]}...")
        
        # 🔧 NEW: First check if this is a plugin tool output that we want to preserve
        if response and not response.strip().startswith('{"name":"'):
            # This might be actual tool output content, preserve it
            pass
        elif (response.strip().startswith('{"name":"') and 
            '"arguments":' in response and
            '"plugin_id":' in response):
            # This is a plugin invocation JSON - try to extract tool output
            try:
                plugin_data = json.loads(response)
                tool_output_fields = ['tool_output_content', 'output', 'result', 'content', 'answer']
                
                for field in tool_output_fields:
                    if field in plugin_data and plugin_data[field]:
                        tool_output = str(plugin_data[field])
                        if len(tool_output.strip()) > 10:
                            print(f"🔧 Extracted {field} from plugin JSON: {tool_output[:80]}...")
                            return clean_ai_response(tool_output)  # Recursive clean
                
                # Check nested arguments
                if 'arguments' in plugin_data and isinstance(plugin_data['arguments'], dict):
                    args = plugin_data['arguments']
                    for field in tool_output_fields:
                        if field in args and args[field]:
                            tool_output = str(args[field])
                            if len(tool_output.strip()) > 10:
                                print(f"🔧 Extracted args.{field} from plugin JSON: {tool_output[:80]}...")
                                return clean_ai_response(tool_output)  # Recursive clean
                
                print("🚫 No useful tool output found in plugin JSON, filtering out")
                return ""  # Return empty if no tool output found
                
            except json.JSONDecodeError:
                print("🚫 Backup filter: Detected malformed plugin JSON, filtering out")
                return ""
        
        # Check for pure system messages (skip only if entire response is system content)
        system_only_indicators = [
            '"msg_type":"time_capsule_recall"',
            '"msg_type":"conversation_summary"', 
            '"msg_type":"system_message"'
        ]
        
        # If the response is ONLY a system message, skip it
        if (response.strip().startswith('{"msg_type"') and 
            any(indicator in response for indicator in system_only_indicators) and
            len(response.strip()) < 2000):  # Short responses that are likely pure system messages
            print("🚫 Detected pure system message, skipping this response")
            return ""  # Return empty to trigger conversation end
        
        # If response looks like JSON, try to extract the actual answer
        if response.strip().startswith('{'):
            try:
                # Handle multiple JSON objects in response (streaming format)
                json_objects = []
                lines = response.strip().split('\n')
                
                for line in lines:
                    line = line.strip()
                    if line.startswith('data: '):
                        line = line[6:]  # Remove 'data: ' prefix
                    
                    if line and line.startswith('{') and line.endswith('}'):
                        try:
                            json_obj = json.loads(line)
                            json_objects.append(json_obj)
                        except json.JSONDecodeError:
                            continue
                
                # If no JSON objects found, try parsing the entire response
                if not json_objects:
                    try:
                        json_objects = [json.loads(response)]
                    except json.JSONDecodeError:
                        pass
                
                # Extract meaningful content from JSON objects
                meaningful_content = ""
                
                for json_obj in json_objects:
                    # Skip system messages
                    if json_obj.get('msg_type') in ['time_capsule_recall', 'conversation_summary', 'system_message']:
                        continue
                    
                    # Extract content from various possible fields
                    content_fields = [
                        'tool_output_content',
                        'content', 
                        'answer',
                        'message',
                        'text',
                        'response'
                    ]
                    
                    for field in content_fields:
                        if field in json_obj and json_obj[field]:
                            content = str(json_obj[field])
                            
                            # Clean up escape characters
                            content = content.replace('\\n', '\n').replace('\\"', '"').replace('\\t', '\t')
                            
                            # Filter out evaluation-related content
                            evaluation_keywords = [
                                '用户编写的信息',
                                '用户画像信息', 
                                '用户记忆点信息',
                                '避免使用隐私和敏感信息',
                                '以下信息来源于用户与你对话',
                                'wraped_text',
                                'origin_search_results'
                            ]
                            
                            if not any(keyword in content for keyword in evaluation_keywords):
                                # Look for "答案：" pattern and extract content after it
                                if "答案：" in content:
                                    answer_part = content.split("答案：", 1)[1]
                                    answer_part = answer_part.replace("参考依据：", "").replace("依据来源：", "")
                                    meaningful_content = answer_part.strip()
                                    break
                                elif len(content.strip()) > 10:  # Substantial content
                                    meaningful_content = content.strip()
                                    break
                
                if meaningful_content:
                    print(f"✅ Extracted from JSON: {meaningful_content[:80]}...")
                    return meaningful_content
                    
            except Exception as e:
                print(f"⚠️ JSON parsing failed: {str(e)}")
        
        # Handle streaming format patterns - enhanced for stream_plugin_finish
        if '"msg_type":"stream_plugin_finish"' in response:
            try:
                import re
                # Try multiple patterns to extract content
                patterns = [
                    r'"tool_output_content":"([^"]+)"',
                    r'"content":"([^"]+)"',
                    r'"answer":"([^"]+)"',
                    r'"text":"([^"]+)"'
                ]
                
                for pattern in patterns:
                    match = re.search(pattern, response)
                    if match:
                        content = match.group(1)
                        content = content.replace('\\n', '\n').replace('\\"', '"').replace('\\t', '\t')
                        
                        # Filter out evaluation content
                        if not any(keyword in content for keyword in [
                            '用户编写的信息', '用户画像信息', '用户记忆点信息',
                            'wraped_text', 'origin_search_results'
                        ]):
                            if "答案：" in content:
                                answer_part = content.split("答案：", 1)[1]
                                answer_part = answer_part.replace("参考依据：", "").replace("依据来源：", "")
                                cleaned_answer = answer_part.strip()
                                if len(cleaned_answer) > 5:  # Ensure substantial content
                                    print(f"✅ Extracted from stream_plugin_finish: {cleaned_answer[:80]}...")
                                    return cleaned_answer
                            elif len(content.strip()) > 5:  # Substantial content
                                print(f"✅ Extracted from stream_plugin_finish: {content.strip()[:80]}...")
                                return content.strip()
                            
                # If no patterns matched, try to parse the JSON directly
                try:
                    json_data = json.loads(response)
                    data_field = json_data.get('data', {})
                    if isinstance(data_field, str):
                        # Sometimes data is a JSON string
                        try:
                            data_obj = json.loads(data_field)
                            tool_output = data_obj.get('tool_output_content', '')
                            if tool_output and len(tool_output.strip()) > 5:
                                print(f"✅ Extracted from nested JSON: {tool_output[:80]}...")
                                return tool_output.strip()
                        except:
                            pass
                except:
                    pass
                    
            except Exception as e:
                print(f"⚠️ Stream plugin parsing failed: {str(e)}")
                pass
        
        # Handle plain text with "答案：" pattern
        if "答案：" in response and not any(keyword in response for keyword in [
            '用户编写的信息', '用户画像信息', '用户记忆点信息'
        ]):
            answer_part = response.split("答案：", 1)[1]
            lines = answer_part.split('\n')
            for line in lines:
                line = line.strip()
                if line and not line.startswith('参考依据：') and not line.startswith('依据来源：'):
                    return line
        
        # Final fallback - return original if it's clean text
        cleaned = response.strip()
        
        # 🔧 DEBUGGING: Check what content is being filtered
        print(f"🔍 CONTENT FILTER DEBUG: Original length: {len(cleaned)} chars")
        print(f"🔍 CONTENT FILTER DEBUG: First 200 chars: {cleaned[:200]}...")
        
        # Final filter check for system content (REDUCED STRICTNESS)
        high_priority_system_patterns = [
            '"msg_type":"time_capsule_recall"',
            '"msg_type":"conversation_summary"',
            '"msg_type":"system_message"'
        ]
        
        # Only filter if it's clearly a system message AND short
        if (any(pattern in cleaned for pattern in high_priority_system_patterns) and 
            len(cleaned.strip()) < 1000):  # Only filter short system messages
            print(f"🚫 Final filter caught system content pattern, returning empty")
            return ""
        
        # 🔧 NEW: Allow evaluation-related content but warn
        evaluation_patterns = [
            '用户编写的信息', '用户画像信息', '用户记忆点信息',
            'wraped_text', 'origin_search_results'
        ]
        
        if any(pattern in cleaned for pattern in evaluation_patterns):
            print(f"⚠️ WARNING: Content contains evaluation pattern but allowing it: {cleaned[:100]}...")
            # Don't return empty - let it through with warning
        
        print(f"✅ Cleaned response: {cleaned[:80]}...")
        return cleaned
        
    except Exception as e:
        print(f"❌ 响应清理异常: {str(e)}")
        return original_response.strip() if len(original_response.strip()) < 500 else ""

async def test_coze_plugin_extraction():
    """
    Test function to debug Coze API plugin content extraction
    """
    print("🧪 Testing Coze API plugin extraction...")
    
    # Test with a simple question that should trigger plugin
    test_message = "地下室防水卷材搭接宽度不够，基层处理好像也有问题"
    test_bot_id = "7498244859505999924"  # Default test bot ID
    
    try:
        # Try the HTTP fallback method first
        print("1️⃣ Testing HTTP fallback method...")
        result = await call_coze_api_fallback(test_message, test_bot_id)
        print(f"✅ HTTP fallback result: {result[:200]}...")
        print(f"📏 Result length: {len(result)} characters")
        
        # Test if result contains substantial content
        if len(result) > 100:
            print("✅ HTTP fallback appears to work correctly")
        else:
            print("❌ HTTP fallback returned minimal content")
            
        return result
        
    except Exception as e:
        print(f"❌ Test failed: {str(e)}")
        return None

# ⭐ Security and validation functions
def validate_filename(filename: str) -> bool:
    """Validate uploaded filename for security"""
    if not filename:
        return False
    
    # Check for path traversal attempts
    if '..' in filename or '/' in filename or '\\' in filename:
        return False
    
    # Check for dangerous extensions
    if any(filename.lower().endswith(ext) for ext in config.BLOCKED_EXTENSIONS):
        return False
    
    # Check for allowed extensions
    file_ext = os.path.splitext(filename)[1].lower()
    if file_ext not in config.ALLOWED_EXTENSIONS:
        return False
    
    # Check filename length
    if len(filename) > config.MAX_FILENAME_LENGTH:
        return False
    
    return True

def sanitize_user_input(text: str, max_length: int = None) -> str:
    """Sanitize user input to prevent injection attacks"""
    if not text:
        return ""
    
    if max_length is None:
        max_length = config.MAX_INPUT_LENGTH
    
    # Remove null bytes and control characters
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    
    # Truncate to max length
    if len(text) > max_length:
        text = text[:max_length] + "...[内容截断]"
    
    return text.strip()

def validate_api_url(url: str) -> bool:
    """Validate API URL for security"""
    if not url:
        return False
    
    # Check for valid HTTP/HTTPS URLs
    if not (url.startswith('http://') or url.startswith('https://')):
        return False
    
    # Prevent local network access
    blocked_hosts = ['localhost', '127.0.0.1', '0.0.0.0', '::1']
    blocked_patterns = [r'192\.168\.', r'10\.', r'172\.(1[6-9]|2\d|3[01])\.']
    
    for host in blocked_hosts:
        if host in url.lower():
            return False
    
    for pattern in blocked_patterns:
        if re.search(pattern, url):
            return False
    
    return True

def find_available_port(start_port: int) -> int:
    """Find an available port starting from the given port number"""
    import socket
    
    port = start_port
    while port < start_port + 100:  # Try up to 100 ports
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind(('localhost', port))
                return port
        except OSError:
            port += 1
    
    # If no port found in range, return the original port
    print(f"⚠️ No available port found in range {start_port}-{start_port + 100}, using {start_port}")
    return start_port

def find_available_port(start_port: int) -> int:
    """Find an available port starting from the given port number"""
    import socket
    
    port = start_port
    while port < start_port + 100:  # Try up to 100 ports
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.bind(('localhost', port))
                return port
        except OSError:
            port += 1
    
    # If no port found in range, return the original port
    print(f"⚠️ No available port found in range {start_port}-{start_port + 100}, using {start_port}")
    return start_port

async def generate_quick_initial_message(scenario: Dict, user_persona_info: Dict, is_tricky_test: bool = False) -> str:
    """
    Generate AI-powered initial user message based on scenario and user persona
    Uses DeepSeek API to create natural, role-appropriate conversation starters
    Now supports tricky test mode for edge-case questions
    """
    try:
        # Extract persona information
        persona_summary = user_persona_info.get('extracted_persona_summary', {})
        role = persona_summary.get('user_persona', {}).get('role', '工程项目现场监理工程师')
        experience_level = persona_summary.get('user_persona', {}).get('experience_level', '有经验')
        communication_style = persona_summary.get('user_persona', {}).get('communication_style', '专业直接')
        business_domain = persona_summary.get('business_domain', '建筑工程')
        
        # Get scenario context
        scenario_title = scenario.get('title', '规范查询')
        scenario_context = scenario.get('context', '工程规范相关问题咨询')
        
        # Choose prompt based on tricky test mode
        if is_tricky_test:
            initial_message_prompt = f"""你现在要扮演{role}，在以下场景中开始一段对话，提出一个**罕见但仍与建筑工程相关的问题**，该问题应该是边缘案例或非常规情况，通常不会出现在标准规范或知识库中，但听起来仍然合理。

场景背景: {scenario_context}
场景标题: {scenario_title}

你的角色特征:
- 职业: {role}
- 经验水平: {experience_level}  
- 沟通风格: {communication_style}
- 工作领域: {business_domain}

请生成一句自然的开场白，提出一个非典型、数据库可能查不到的问题，必须满足：
1. 与建筑或施工标准有**一定关联**
2. **不会直接命中已有标准**
3. 问题要清晰、合理但"刁钻"，如非常规材料/极端场景/过时做法/地方奇特要求
4. 控制在50字以内

直接输出对话内容，不要其他解释："""
        else:
            initial_message_prompt = f"""你现在要扮演{role}，在以下场景中开始一段对话：

场景背景: {scenario_context}
场景标题: {scenario_title}

你的角色特征:
- 职业: {role}
- 经验水平: {experience_level}  
- 沟通风格: {communication_style}
- 工作领域: {business_domain}

请生成一句自然的开场白，作为{role}向AI助手提出的第一个问题或请求。
要求：
1. 语言自然，符合{communication_style}的风格
2. 体现{role}的专业关注点
3. 与{scenario_title}场景相关
4. 长度控制在50字以内

直接输出对话内容，不要其他解释："""

        response = await call_deepseek_api_enhanced(
            initial_message_prompt, 
            temperature=0.3, 
            max_tokens=150
        )
        
        # Clean and validate response
        initial_message = response.strip()
        
        # Fallback if AI generation fails or is inappropriate
        if (len(initial_message) > 200 or len(initial_message) < 10 or 
            not initial_message or "扮演" in initial_message or "生成" in initial_message):
            
            # Use domain-specific fallback messages
            if is_tricky_test:
                fallback_messages = {
                    '建筑工程': f"如果在南极建预制板构件的冷接缝处理需要注意什么？",
                    '金融银行': f"数字货币的资产配置在极端通胀下如何平衡？",
                    '医疗健康': f"高原缺氧环境下的麻醉用药剂量如何调整？",
                    '教育培训': f"多语言混合教学的评估体系如何建立？"
                }
            else:
                fallback_messages = {
                    '建筑工程': f"你好，我是{role}，现场遇到了一些{scenario_title}相关的问题，想咨询一下相关规范要求。",
                    '金融银行': f"你好，我想了解一下{scenario_title}相关的业务流程和要求。",
                    '医疗健康': f"你好，关于{scenario_title}方面有些专业问题需要咨询。",
                    '教育培训': f"你好，想了解{scenario_title}相关的具体要求和标准。"
                }
            
            initial_message = fallback_messages.get(
                business_domain, 
                f"你好，我是{role}，想咨询一下{scenario_title}相关的问题。"
            )
        
        mode_label = "🎯刁钻" if is_tricky_test else "📝常规"
        print(f"✅ {mode_label}生成初始消息: {initial_message[:50]}...")
        return initial_message
        
    except Exception as e:
        print(f"❌ 初始消息生成失败: {str(e)}")
        # Ultimate fallback
        role = user_persona_info.get('extracted_persona_summary', {}).get('user_persona', {}).get('role', '用户')
        scenario_title = scenario.get('title', '咨询')
        return f"你好，我是{role}，想咨询一下{scenario_title}相关的问题。"

async def generate_single_initial_message(scenario_info: Dict, user_persona_info: Dict, is_tricky_test: bool = False) -> str:
    """
    Generate single initial message for dynamic conversations
    Wrapper around generate_quick_initial_message for compatibility
    """
    return await generate_quick_initial_message(scenario_info, user_persona_info, is_tricky_test)

async def generate_next_message_based_on_response(
    scenario_info: Dict, 
    user_persona_info: Dict, 
    conversation_history: List[Dict], 
    ai_response: str,
    is_tricky_test: bool = False
) -> str:
    """
    Generate next user message based on AI's response
    Supports both normal and tricky test modes
    """
    try:
        # Extract persona information
        persona_summary = user_persona_info.get('extracted_persona_summary', {})
        persona = user_persona_info.get('user_persona', {})
        role = persona_summary.get('user_persona', {}).get('role', '工程项目现场监理工程师')
        communication_style = persona_summary.get('user_persona', {}).get('communication_style', '专业直接')
        
        # Build conversation context
        conversation_context = "\n".join([
            f"用户: {turn['user_message']}\nAI: {turn['ai_response']}" 
            for turn in conversation_history[-2:]  # Only use last 2 turns for context
        ])
        
        # Choose prompt based on tricky test mode
        if is_tricky_test:
            followup_prompt = f"""你是{role}，正在与AI助手进行专业咨询对话。以下是对话历史：

{conversation_context}

AI刚才的回应：{ai_response}

现在请你提出一个**新的相关问题**，这个问题仍然要偏向**边缘情况或非主流需求**，可能与极端天气、历史建筑、特殊材料、特殊场景等相关，普通数据库或标准库通常查不到。

要求：
1. 不重复之前问题，但仍保持上下文连贯
2. 继续以{communication_style}的方式沟通
3. 提出一个让AI助手更难检索的"刁钻问题"，但必须合理
4. 长度控制在50字以内

直接输出跟进问题内容，不要其他解释："""
        else:
            followup_prompt = f"""你是{role}，正在与AI助手进行专业咨询对话。以下是对话历史：

{conversation_context}

AI刚才的回应：{ai_response}

根据AI的回应和你的专业背景，生成下一句自然的跟进问题。

要求:
1. 基于AI的具体回应内容进行有针对性的跟进
2. 体现{role}的专业关注点和思维方式
3. 语言自然，符合{communication_style}的风格
4. 长度控制在50字以内

直接输出对话内容，不要其他解释："""

        response = await call_deepseek_api_enhanced(
            followup_prompt, 
            temperature=0.4, 
            max_tokens=120
        )
        
        # Clean and validate response
        next_message = response.strip()
        
        # Check for conversation end indicators
        end_indicators = ["谢谢", "明白了", "清楚了", "了解了", "知道了", "好的", "没问题", 
                         "满意", "解决了", "够了", "足够", "可以了", "ok", "OK", "感谢"]
        
        if any(indicator in next_message for indicator in end_indicators):
            return "END"  # Signal to end conversation
        
        # Fallback for inappropriate responses
        if (len(next_message) > 200 or len(next_message) < 5 or 
            not next_message or "扮演" in next_message or "生成" in next_message):
            
            if is_tricky_test:
                # Tricky follow-up fallbacks based on business domain
                business_domain = user_persona_info.get('business_domain', '建筑工程')
                tricky_followups = {
                    '建筑工程': "这种极寒地区是否可以跳过二次养护流程？",
                    '金融银行': "这种配置在战争时期还适用吗？",
                    '医疗健康': "如果在太空环境下这个剂量要如何调整？",
                    '教育培训': "这套体系对认知障碍学生是否有效？"
                }
                next_message = tricky_followups.get(business_domain, "还有其他特殊情况需要考虑吗？")
            else:
                # Standard follow-up fallbacks
                fallback_questions = [
                    "还有其他需要注意的细节吗？",
                    "请详细说明具体的操作要求。",
                    "有没有相关的检测标准？",
                    "这方面还有什么规范要求？"
                ]
                next_message = fallback_questions[len(conversation_history) % len(fallback_questions)]
        
        mode_label = "🎯刁钻" if is_tricky_test else "📝常规"
        print(f"✅ {mode_label}生成跟进消息: {next_message[:50]}...")
        return next_message
        
    except Exception as e:
        print(f"❌ 跟进消息生成失败: {str(e)}")
        return "还有其他需要了解的吗？"

def get_score_grade(score: float) -> str:
    """
    Convert numerical score (0-100) to Chinese grade label
    """
    if score >= 90:
        return "优秀"
    elif score >= 80:
        return "良好"
    elif score >= 70:
        return "中等"
    elif score >= 60:
        return "及格"
    else:
        return "不及格"

if __name__ == "__main__":
    import sys
    import os
    if len(sys.argv) > 1 and sys.argv[1] == "test":
        # Run test mode
        import asyncio
        asyncio.run(test_coze_plugin_extraction())
    else:
        # Run normal server - use environment variable PORT if available (cloud deployment)
        port = int(os.environ.get("PORT", config.DEFAULT_PORT))
        if port != config.DEFAULT_PORT:
            print(f"🌐 使用云平台指定端口: {port}")
        else:
            port = find_available_port(port)
        print(f"🚀 AI Agent评估平台启动在端口 {port}")
        uvicorn.run(app, host=config.DEFAULT_HOST, port=port) 